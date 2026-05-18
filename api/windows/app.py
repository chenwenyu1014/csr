#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Windows Bridge Service (FastAPI)

端点：
- GET  /healthz                                  健康检查
- GET  /version                                  版本信息
- POST /api/v1/rtf/insert_head_section_break     为RTF在文首插入“下一页分节符”，返回处理后的RTF

实现策略（优先级）：
1) Spire.Doc：加载RTF -> 文首插入空段落 -> 该段后插入 SectionBreakType.NewPage -> 返回RTF
2) Word COM：打开 -> 在Range(0,0)插入 wdSectionBreakNextPage(2) -> 保存RTF -> 返回
3) 失败回退：返回原始RTF
"""

from __future__ import annotations  # 兼容前向引用的类型注解

import os  # 标准库：环境变量与路径
import sys  # 标准库：系统路径
import tempfile  # 标准库：临时目录/文件
import shutil  # 标准库：文件复制
import threading
from contextlib import contextmanager

# 无论从哪里启动，都能正确导入
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)
import uuid  # 标准库：请求ID/随机名
import logging  # 标准库：日志记录
import time  # 标准库：耗时统计
import json  # 标准库：JSON 读写
from typing import Optional, Dict, Any
from fastapi import FastAPI, Request, HTTPException, Form
from fastapi.responses import JSONResponse
from pathlib import Path
from dotenv import load_dotenv, find_dotenv

# 预加载环境变量，确保 setup_logging 能读取到 ENVIRONMENT 配置
try:
    _env_path = find_dotenv(usecwd=True)
    if not _env_path:
        _env_path = str(Path(__file__).resolve().parents[2] / ".env")
    load_dotenv(_env_path, override=False, encoding="utf-8")
except Exception:
    pass

# 使用统一日志配置（控制台 + 文件双输出）
from utils.logging_config import setup_logging
setup_logging(service_name="Bridge")

# 设置本模块的logger
logger = logging.getLogger(__name__)
app = FastAPI(title="Windows Bridge Service", version="1.0.0")
# ============================================================

_WIN_BRIDGE_TASK_LOCK = threading.Lock()


def _get_request_id_from_request(request: Request) -> str:
    try:
        rid = request.headers.get("X-Request-Id") or request.headers.get("x-request-id")
        return (rid or "-").strip() or "-"
    except Exception:
        return "-"


@contextmanager
def _windows_serial_task_guard(task_name: str, request: Request):
    """
    串行化 Windows 端重任务（尤其是 Word/Excel COM）。
    - wait 模式：排队等待锁
    - reject 模式：锁被占用则直接返回 429
    """
    mode = (os.getenv("WINDOWS_BRIDGE_SERIAL_MODE", "wait") or "wait").strip().lower()
    timeout_raw = (os.getenv("WINDOWS_BRIDGE_SERIAL_TIMEOUT", "0") or "0").strip()
    try:
        timeout_s = float(timeout_raw)
    except Exception:
        timeout_s = 0.0

    rid = _get_request_id_from_request(request)

    # 1) 线程级互斥（同进程内串行）
    started_wait = time.perf_counter()
    acquired = False
    if mode == "reject":
        acquired = _WIN_BRIDGE_TASK_LOCK.acquire(blocking=False)
    else:
        if timeout_s and timeout_s > 0:
            acquired = _WIN_BRIDGE_TASK_LOCK.acquire(timeout=timeout_s)
        else:
            _WIN_BRIDGE_TASK_LOCK.acquire()
            acquired = True

    if not acquired:
        raise HTTPException(
            status_code=429,
            detail=f"windows-bridge busy: another task is running (task={task_name}, rid={rid})",
        )

    wait_ms = int((time.perf_counter() - started_wait) * 1000)
    if wait_ms >= 50:
        logger.info(f"⏳ 等待Windows串行锁 {wait_ms}ms (task={task_name}, rid={rid})")

    # 2) 可选：跨进程文件锁（防止误用多 worker / 多实例）
    ipc_enabled = (os.getenv("WINDOWS_BRIDGE_SERIAL_IPC_LOCK", "1") or "1").strip()
    lock_f = None
    try:
        if ipc_enabled not in ("0", "false", "False", "no", "NO"):
            lock_path = (os.getenv("WINDOWS_BRIDGE_SERIAL_LOCK_FILE", "AAA/.windows_bridge.lock") or "").strip()
            if not lock_path:
                lock_path = "AAA/.windows_bridge.lock"
            lp = Path(lock_path)
            try:
                lp.parent.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass

            try:
                import msvcrt  # Windows only
                lock_f = open(lp, "a+b")
                try:
                    # 确保文件至少有 1 字节，并将锁定区间固定在 offset=0
                    lock_f.seek(0, os.SEEK_END)
                    if lock_f.tell() == 0:
                        lock_f.write(b"\0")
                        lock_f.flush()
                    lock_f.seek(0)
                except Exception:
                    try:
                        lock_f.seek(0)
                    except Exception:
                        pass
                # 锁 1 字节区间；阻塞等待即可（线程锁已经串行，跨进程才会竞争）
                msvcrt.locking(lock_f.fileno(), msvcrt.LK_LOCK, 1)
            except Exception:
                # 文件锁为 best-effort：失败不阻断业务，但会失去跨进程保护
                try:
                    if lock_f:
                        lock_f.close()
                except Exception:
                    pass
                lock_f = None

        yield
    finally:
        # 释放文件锁
        if lock_f is not None:
            try:
                import msvcrt  # type: ignore
                try:
                    lock_f.seek(0)
                except Exception:
                    pass
                try:
                    msvcrt.locking(lock_f.fileno(), msvcrt.LK_UNLCK, 1)
                except Exception:
                    pass
            finally:
                try:
                    lock_f.close()
                except Exception:
                    pass
        # 释放线程锁
        try:
            _WIN_BRIDGE_TASK_LOCK.release()
        except Exception:
            pass

# ========== Content Control插入器（完整版，支持横竖方向检测和分节符）==========

# Word常量
wdCollapseEnd = 0
wdCollapseStart = 1
wdSectionBreakNextPage = 2
wdOrientLandscape = 1
wdOrientPortrait = 0


# ========== 导入核心插入模块 ==========

from pathlib import Path
from service.windows.insertion.word_control_content_inserter import WordControlContentInserter

# 异步任务管理
task_storage: Dict[str, Dict[str, Any]] = {}
task_results: Dict[str, Any] = {}



# 统一 JSON 日志
# setup_json_logging 暂时不使用，直接用标准logging



def _auth_ok(request: Request) -> bool:
    token = (os.getenv("WINDOWS_BRIDGE_TOKEN") or "").strip()
    if not token:
        return True
    auth = request.headers.get("Authorization") or request.headers.get("authorization") or ""
    if not auth.lower().startswith("bearer "):
        return False
    return auth.split(" ", 1)[1].strip() == token


# 请求链路ID上下文（独立于主服务）
try:
    from contextvars import ContextVar
    request_id_ctx = ContextVar("request_id", default="-")
except Exception:
    request_id_ctx = None  # type: ignore


@app.middleware("http")
async def _with_request_id(request: Request, call_next):
    rid = request.headers.get("X-Request-Id") or ("req_" + uuid.uuid4().hex)
    token = None
    if request_id_ctx is not None:
        try:
            token = request_id_ctx.set(rid)
        except Exception:
            token = None
    started = time.perf_counter()
    response = None
    try:
        response = await call_next(request)
        return response
    finally:
        try:
            dur_ms = int((time.perf_counter() - started) * 1000)
            status = getattr(response, "status_code", 0) if response is not None else 0
            ua = request.headers.get("user-agent") or request.headers.get("User-Agent")
            ref = request.headers.get("referer") or request.headers.get("Referer")
            clen = request.headers.get("content-length") or request.headers.get("Content-Length")
            try:
                clen_val = int(clen) if clen else None
            except Exception:
                clen_val = None
            route_path = None
            try:
                route = request.scope.get("route")
                route_path = getattr(route, "path", None)
            except Exception:
                route_path = None
            logging.getLogger("bridge.access").info(
                "request.done",
                extra={
                    "event": "request.done",
                    "path": request.url.path if hasattr(request, "url") else None,
                    "route": route_path,
                    "method": getattr(request, "method", None),
                    "status": status,
                    "duration_ms": dur_ms,
                    "client": getattr(getattr(request, "client", None), "host", None),
                    "remote_port": getattr(getattr(request, "client", None), "port", None),
                    "user_agent": ua,
                    "referer": ref,
                    "request_size": clen_val,
                }
            )
        except Exception:
            pass
        try:
            if response is not None:
                response.headers["X-Request-Id"] = rid
        except Exception:
            pass
        if token is not None:
            try:
                request_id_ctx.reset(token)  # type: ignore
            except Exception:
                pass

@app.get("/healthz")
def healthz() -> JSONResponse:
    info = {
        "spire_available": _probe_spire_available(),
        "win32_available": _probe_win32_available(),
    }
    return JSONResponse({"status": "ok", "info": info})


@app.get("/version")
def version() -> JSONResponse:
    return JSONResponse({"service": "windows-bridge", "version": "1.0.0"})


def _probe_spire_available() -> bool:
    try:
        import spire.doc  # type: ignore
        return True
    except Exception:
        return False


def _probe_win32_available() -> bool:
    try:
        import win32com.client  # type: ignore
        return True
    except Exception:
        return False


# ---------- 清理文档（python-docx，优先方案） ----------

def _clean_document_with_docx(file_path: str, output_path: str, remove_first_line: bool = True) -> Dict[str, Any]:
    """
    用 python-docx 清理文档（不依赖 COM，速度快）

    功能：
    1. 删除首行（水印）
    2. 清理 Content Control 控件（保留内容）
    3. 删除占位符标记段落

    Args:
        file_path: 输入文件路径
        output_path: 输出文件路径
        remove_first_line: 是否删除首行

    Returns:
        Dict: 包含成功状态、清理信息等
    """
    result = {
        "success": False,
        "controls_removed": 0,
        "first_line_removed": False,
        "markers_removed": False,
        "error": None
    }

    try:
        from docx import Document
        import re

        doc = Document(file_path)

        # 1. 删除首行
        if remove_first_line and len(doc.paragraphs) > 0:
            first_para = doc.paragraphs[0]._element
            parent = first_para.getparent()
            if parent is not None:
                parent.remove(first_para)
                result["first_line_removed"] = True
                logger.info("✅ 已删除首行")

        # 2. 清理 Content Control（只删除有 tag 的自定义控件，保留内容）
        body = doc._element.body
        WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # 找到所有 sdt 元素
        sdt_elements = body.findall(f'.//{{{WNS}}}sdt')
        controls_count = 0

        for sdt in sdt_elements:
            try:
                # 检查是否有 tag（自定义控件有 tag，目录控件等没有）
                sdt_pr = sdt.find(f'{{{WNS}}}sdtPr')
                if sdt_pr is not None:
                    tag_elem = sdt_pr.find(f'{{{WNS}}}tag')
                    if tag_elem is None:
                        # 没有 tag，不是自定义控件（如目录），跳过
                        continue

                # 获取 sdtContent 里的内容
                sdt_content = sdt.find(f'{{{WNS}}}sdtContent')
                if sdt_content is not None:
                    # 取出 sdtContent 的所有子元素
                    children = list(sdt_content)
                    parent = sdt.getparent()
                    if parent is not None:
                        # 用 addnext 从后往前插入，保证最终顺序正确
                        for child in reversed(children):
                            sdt.addnext(child)
                        parent.remove(sdt)
                        controls_count += 1
            except Exception:
                pass

        result["controls_removed"] = controls_count
        logger.info(f"✅ 已清理 {controls_count} 个 Content Control 控件")

        # 3. 删除占位符标记段落
        marker_regex = re.compile(r'^\{\{(TemplateTable_\d+_Start|TemplateTable_\d+_End|标题：[^}]+)\}\}$')

        to_delete = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if marker_regex.match(text):
                to_delete.append(para._element)

        deleted_markers = 0
        for elem in to_delete:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                deleted_markers += 1

        result["markers_removed"] = deleted_markers > 0
        if deleted_markers > 0:
            logger.info(f"✅ 已清理 {deleted_markers} 个占位符标记段落")
        else:
            logger.info("ℹ️ 文档中没有占位符标记")

        # 保存
        doc.save(output_path)
        result["success"] = True
        result["output_file"] = output_path

    except Exception as e:
        import traceback
        traceback.print_exc()
        result["error"] = str(e)
        logger.error(f"python-docx 清理失败: {e}", exc_info=True)

    return result


# ---------- 清理Content Control并保留内容（COM，回退方案） ----------

def _clean_content_controls_preserve_content(file_path: str, output_path: str, remove_first_line: bool = True) -> Dict[str, Any]:
    """
    清理Word文档中的Content Control控件，但保留控件内的内容
    
    Args:
        file_path: 输入文件路径
        output_path: 输出文件路径
        remove_first_line: 是否删除首行（水印）
    
    Returns:
        Dict: 包含成功状态、清理的控件数量等信息
    """
    try:
        import win32com.client as win32  # type: ignore
    except Exception:
        return {"success": False, "error": "win32com不可用"}
    from utils.windows_com import safe_dispatch
    
    result = {
        "success": False,
        "controls_removed": 0,
        "first_line_removed": False,
        "markers_removed": False,
        "error": None
    }
    
    word = None
    doc = None
    
    com_inited = False
    try:
        # 初始化COM
        import pythoncom
        try:
            pythoncom.CoInitialize()
            com_inited = True
        except:
            pass
        
        word = safe_dispatch("Word.Application", use_ex=False, logger=logger)
        try:
            word.Visible = False
            word.DisplayAlerts = 0
        except Exception:
            pass
        
        # 打开文档
        doc = word.Documents.Open(str(Path(file_path).resolve()), ReadOnly=False)
        logger.info(f"📄 打开文档: {file_path}")
        
        # 1. 清理Content Control控件（保留内容）
        controls_count = 0
        try:
            # 获取所有Content Control
            content_controls = doc.ContentControls
            total_controls = content_controls.Count
            logger.info(f"📝 发现 {total_controls} 个Content Control控件")
            
            # 从后往前删除，避免索引问题
            # 注意：COM集合是1-indexed
            for i in range(total_controls, 0, -1):
                try:
                    cc = content_controls.Item(i)
                    # Delete(False) = 删除控件但保留内容
                    cc.Delete(False)
                    controls_count += 1
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    logger.warning(f"删除控件 {i} 失败: {e}")
            
            result["controls_removed"] = controls_count
            if total_controls == 0:
                logger.info("ℹ️ 文档中没有Content Control控件")
                result["controls_message"] = "文档中没有Content Control控件"
            else:
                logger.info(f"✅ 已清理 {controls_count}/{total_controls} 个Content Control控件")
        except Exception as e:
            import traceback
            traceback.print_exc()
            logger.warning(f"清理Content Control失败: {e}")
        
        # 2. 删除首行（水印）
        if remove_first_line:
            try:
                para = None
                try:
                    # 尝试获取第一节的第一段
                    para = doc.Sections(1).Range.Paragraphs(1)
                except Exception:
                    # 回退到文档的第一段
                    para = doc.Paragraphs(1)
                
                if para is not None:
                    rng = para.Range
                    # 获取首行文本用于日志
                    first_line_text = rng.Text[:50] if rng.Text else ""
                    logger.info(f"📝 首行内容: {repr(first_line_text)}...")
                    
                    # 删除整个段落（包括段落标记）
                    rng.Delete()
                    result["first_line_removed"] = True
                    logger.info("✅ 已删除首行")
            except Exception as e:
                import traceback
                traceback.print_exc()
                logger.warning(f"删除首行失败: {e}")

        # 3. 清理模板占位符标记 {{Table_*_Start}}、{{Table_*_End}}、{{标题：...}}
        # 用 python-docx 处理（更快更可靠）
        result["markers_removed"] = False
        try:
            from docx import Document
            import re

            # 标记匹配正则
            marker_regex = re.compile(r'^\{\{(Table_\d+_Start|Table_\d+_End|标题：[^}]+)\}\}$')

            # 先保存当前 COM 文档状态
            doc.Save()
            doc.Close(SaveChanges=False)

            # 用 python-docx 打开并处理
            docx_doc = Document(str(Path(output_path).resolve()))

            # 收集要删除的段落元素
            to_delete = []
            for para in docx_doc.paragraphs:
                text = para.text.strip()
                if marker_regex.match(text):
                    to_delete.append(para._element)

            # 删除段落元素（直接操作 XML）
            deleted_count = 0
            for elem in to_delete:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
                    deleted_count += 1

            # 保存
            docx_doc.save(str(Path(output_path).resolve()))

            result["markers_removed"] = deleted_count > 0

            if deleted_count > 0:
                logger.info(f"✅ 已清理 {deleted_count} 个占位符标记段落")
            else:
                logger.info("ℹ️ 文档中没有占位符标记")

            # 重新打开文档（后续流程需要）
            doc = word.Documents.Open(str(Path(output_path).resolve()), ReadOnly=False)

        except Exception as e:
            import traceback
            traceback.print_exc()
            logger.warning(f"清理占位符标记失败: {e}")

        # 保存到输出路径
        output_path_resolved = str(Path(output_path).resolve())
        # 确保输出目录存在
        Path(output_path_resolved).parent.mkdir(parents=True, exist_ok=True)
        
        # 保存文件
        doc.SaveAs(output_path_resolved)
        logger.info(f"💾 已保存到: {output_path}")
        
        result["success"] = True
        result["output_file"] = output_path
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.error(f"清理文档失败: {e}", exc_info=True)
        result["error"] = str(e)
    
    finally:
        # 清理资源
        try:
            if doc is not None:
                doc.Close(SaveChanges=False)  # 已经手动保存了
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            if com_inited:
                import pythoncom  # type: ignore
                pythoncom.CoUninitialize()
        except Exception:
            pass
    
    return result


# ---------- 清理首行文字（COM） ----------

def _clear_first_line_with_com_bytes(src_bytes: bytes, suffix: str) -> Optional[bytes]:
    try:
        import win32com.client as win32  # type: ignore
    except Exception:
        return None
    from utils.windows_com import safe_dispatch
    try:
        with tempfile.TemporaryDirectory() as td:
            td = str(td)
            inp = os.path.join(td, f"in_{uuid.uuid4().hex[:8]}{suffix}")
            outp = os.path.join(td, f"out_{uuid.uuid4().hex[:8]}{suffix}")
            with open(inp, "wb") as f:
                f.write(src_bytes)

            com_inited = False
            try:
                import pythoncom  # type: ignore
                try:
                    pythoncom.CoInitialize()
                    com_inited = True
                except Exception:
                    pass
            except Exception:
                com_inited = False

            word = safe_dispatch("Word.Application", use_ex=False, logger=logger)
            try:
                word.Visible = False
                word.DisplayAlerts = 0
            except Exception:
                pass
            doc = None
            try:
                doc = word.Documents.Open(inp, ReadOnly=False)
                try:
                    para = None
                    try:
                        para = doc.Sections(1).Range.Paragraphs(1)
                    except Exception:
                        para = doc.Paragraphs(1)
                    if para is not None:
                        rng = para.Range
                        try:
                            rng.End = rng.End - 1
                        except Exception:
                            pass
                        rng.Text = ""
                except Exception:
                    pass
                doc.SaveAs(outp)
                data = open(outp, "rb").read()
            finally:
                try:
                    if doc is not None:
                        doc.Close(SaveChanges=True)
                except Exception:
                    pass
                try:
                    word.Quit()
                except Exception:
                    pass
                try:
                    if com_inited:
                        import pythoncom  # type: ignore
                        pythoncom.CoUninitialize()
                except Exception:
                    pass
            return data
    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.warning(f"COM 清理首行失败: {e}")
        return None


@app.post("/ky/sys/ai/insert_direct")
def content_control_insert_direct(
    request: Request,
    template_file: str = Form(...),  # 模板文件路径（相对于AAA）
    data_json: str = Form(...)       # JSON数据
):
    """
    直接从共享文件夹插入内容（无需zip）
    
    注意：此函数故意使用同步def而非async def，因为内部有阻塞的COM操作。
    FastAPI会自动在线程池中执行同步函数，避免阻塞事件循环。
    
    Args:
        template_file: 模板文件路径，相对于AAA目录
        data_json: JSON字符串，包含generation_results和resource_mappings
    """
    if not _auth_ok(request):
        raise HTTPException(status_code=401, detail="unauthorized")
    
    with _windows_serial_task_guard("insert_direct", request):
        return _content_control_insert_direct_impl(template_file=template_file, data_json=data_json, request=request)


def _content_control_insert_direct_impl(*, template_file: str, data_json: str, request: Request):
    logger.info("=" * 70)
    logger.info("Content Control插入服务（直接模式）")
    logger.info("=" * 70)
    
    try:
        # ✅ 显示原始参数（用于调试）
        logger.info(f"📄 接收到的template_file: {template_file}")
        logger.info(f"📄 接收到的data_json长度: {len(data_json)} 字符")
        # logger.info(f"📄 data_json前500字符: {data_json[:500]}")
        
        # 解析JSON数据
        data = json.loads(data_json)
        
        # ✅ 兼容两种字段名：generation_results（新）和 paragraphs（旧）
        generation_results = data.get('generation_results') or data.get('paragraphs', [])
        
        # ✅ 标准化字段名：paragraph_id -> control_title
        #    同时确保有generated_content字段
        for item in generation_results:
            if 'paragraph_id' in item and 'control_title' not in item:
                item['control_title'] = item['paragraph_id']
            # 确保有generated_content字段（兼容content字段）
            if 'generated_content' not in item and 'content' in item:
                item['generated_content'] = item['content']
            # 确保有status字段
            if 'status' not in item:
                item['status'] = 'success'

            # 兼容 generated_content 为 dict 对象的情况（表格 JSON）
            # 如果是 dict 类型，转为 JSON 字符串，以便下游 _is_table_json() 检测
            gc = item.get('generated_content')
            if isinstance(gc, dict):
                item['generated_content'] = json.dumps(gc, ensure_ascii=False)
                logger.info(f"  🔄 generated_content 从 dict 转为 JSON 字符串（段落: {item.get('control_title')}）")

            # 清理段落级 resource_mappings 的路径
            paragraph_mappings = item.get('resource_mappings', {})
            if paragraph_mappings:
                cleaned_mappings = {}
                for placeholder, mapping in paragraph_mappings.items():
                    rel_path = mapping.get('path', '') if isinstance(mapping, dict) else str(mapping)

                    # 统一路径分隔符
                    clean_path = rel_path.replace("\\", "/")

                    # 处理混合路径格式（如 /home/xxx/AAA/project_data/...）
                    aaa_idx = clean_path.lower().find("/aaa/")
                    if aaa_idx != -1:
                        clean_path = clean_path[aaa_idx + 5:]  # 跳过 /AAA/
                    elif clean_path.lower().startswith("aaa/"):
                        clean_path = clean_path[4:]  # 跳过 AAA/
                    elif clean_path.startswith("/"):
                        clean_path = clean_path[1:]  # 去掉开头的 /

                    # 转换为本地相对路径
                    resource_path = f"../AAA/{clean_path}"
                    if not Path(resource_path).exists():
                        resource_path = f"AAA/{clean_path}"

                    # 保留原有映射结构
                    if isinstance(mapping, dict):
                        cleaned_mappings[placeholder] = {**mapping, 'path': resource_path}
                    else:
                        cleaned_mappings[placeholder] = {'path': resource_path}

                    logger.info(f"  段落资源路径清理: {placeholder} -> {Path(resource_path).name}")

                item['resource_mappings'] = cleaned_mappings

        # ✅ 清理模板路径：确保是相对路径
        # 先去除首尾的空白和引号
        clean_template = template_file.strip().strip('"').strip("'")
        
        # 处理各种路径前缀
        if clean_template.startswith('//'):
            # 处理 //project_data/... 格式，转换为 project_data/...
            clean_template = clean_template[2:]
        elif clean_template.startswith('/AAA/'):
            clean_template = clean_template[5:]
        elif clean_template.startswith('AAA/'):
            clean_template = clean_template[4:]
        elif clean_template.startswith('/'):
            clean_template = clean_template[1:]
        
        # 记录清理后的路径
        logger.info(f"📁 清理后的模板路径: {clean_template}")
        
        # 直接使用相对路径（相对于AAA）
        # 从父目录查找AAA（因为在windows_bridge目录下运行）
        template_path = f"../AAA/{clean_template}"
        logger.info(f"📁 尝试路径1: {template_path} (存在: {Path(template_path).exists()})")
        
        if not Path(template_path).exists():
            # 也尝试当前目录
            template_path = f"AAA/{clean_template}"
            logger.info(f"📁 尝试路径2: {template_path} (存在: {Path(template_path).exists()})")
            
            if not Path(template_path).exists():
                # 也尝试直接路径（如果已经是完整路径）
                template_path_direct = clean_template
                logger.info(f"📁 尝试路径3: {template_path_direct} (存在: {Path(template_path_direct).exists()})")
                
                if not Path(template_path_direct).exists():
                    error_msg = f"模板文件不存在。尝试的路径:\n  1. ../AAA/{clean_template}\n  2. AAA/{clean_template}\n  3. {clean_template}\n原始template_file参数: {template_file}"
                    logger.error(error_msg)
                    raise FileNotFoundError(error_msg)
                else:
                    template_path = template_path_direct
        
        logger.info(f"模板文件: {template_path}")
        logger.info(f"段落数: {len(generation_results)}")
        
        # 输出文件路径（相对路径）
        output_dir = "../AAA/output"
        if not Path("../AAA").exists():
            output_dir = "AAA/output"
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        import time
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        output_file = f"{output_dir}/result_{timestamp}.docx"
        
        logger.info(f"输出文件: {output_file}")
        
        # 执行插入
        inserter = WordControlContentInserter()
        result = inserter.insert_to_template(
            template_file=template_path,
            generation_results=generation_results,
            output_file=output_file
        )
        
        if result.success:
            logger.info("✅ 插入成功")
            logger.info(f"   - 插入控件: {len(result.inserted_controls)} 个")
            logger.info(f"   - 插入资源: {len(result.inserted_resources)} 个")
            logger.info(f"   - 输出: {output_file}")
            
            # 返回文件路径
            return JSONResponse({
                "success": True,
                "output_file": output_file,
                "inserted_controls": len(result.inserted_controls),
                "inserted_resources": len(result.inserted_resources)
            })
        else:
            raise HTTPException(status_code=500, detail=f"插入失败: {result.error}")

    except HTTPException as e:
        logger.error(f"❌ 处理失败: {e.detail}")
        raise e
    except FileNotFoundError as e:
        error_msg = f"文件不存在: {str(e)}"
        logger.error(f"❌ {error_msg}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=error_msg)
    except Exception as e:
        import traceback
        traceback.print_exc()
        error_msg = str(e) if str(e) else f"未知错误: {type(e).__name__}"
        logger.error(f"❌ 处理失败: {error_msg}")
        import traceback
        tb_str = traceback.format_exc()
        logger.error(f"❌ 堆栈:\n{tb_str}")
        raise HTTPException(status_code=500, detail=error_msg)



# ========== 文档清理接口 ==========

@app.post("/api/v1/document/clean")
def clean_document(
    request: Request,
    file_path: str = Form(..., description="文件路径（相对于AAA目录）"),
    output_path: str = Form(None, description="输出文件路径（可选，默认覆盖原文件）"),
    remove_first_line: bool = Form(True, description="是否删除首行（水印）"),
    remove_content_controls: bool = Form(True, description="是否清理Content Control控件"),
):
    """
    清理Word文档接口（同步执行，避免阻塞事件循环）
    
    功能：
    1. 清理Content Control控件（保留控件内的内容）
    2. 删除文件首行（通常是水印）
    
    Args:
        file_path: 文件路径，相对于AAA目录
        output_path: 输出文件路径（可选），如果不提供则覆盖原文件
        remove_first_line: 是否删除首行（默认True）
        remove_content_controls: 是否清理Content Control（默认True）
    
    Returns:
        JSON响应，包含清理结果
    """
    if not _auth_ok(request):
        raise HTTPException(status_code=401, detail="unauthorized")

    with _windows_serial_task_guard("document_clean", request):
        return _clean_document_impl(
            request=request,
            file_path=file_path,
            output_path=output_path,
            remove_first_line=remove_first_line,
            remove_content_controls=remove_content_controls,
        )


def _clean_document_impl(
    *,
    request: Request,
    file_path: str,
    output_path: str | None,
    remove_first_line: bool,
    remove_content_controls: bool,
):
    logger.info("=" * 70)
    logger.info("文档清理服务")
    logger.info("=" * 70)
    
    try:
        # 清理文件路径（去除引号、空格、统一分隔符）
        clean_path = file_path.strip().strip('"').strip("'")
        clean_path = clean_path.replace("\\", "/")  # 统一使用正斜杠
        
        logger.info(f"📥 原始路径: {repr(file_path)}")
        logger.info(f"📥 清理后路径: {clean_path}")
        
        # 去除AAA前缀
        if clean_path.startswith('/AAA/'):
            clean_path = clean_path[5:]
        elif clean_path.startswith('AAA/'):
            clean_path = clean_path[4:]
        elif clean_path.startswith('/'):
            clean_path = clean_path[1:]
        
        logger.info(f"📥 相对路径: {clean_path}")
        
        # 构建完整路径
        AAA_ROOT = Path(os.getenv("WINDOWS_AAA_ROOT", "AAA")).resolve()
        
        # 尝试多个可能的基础路径
        base_paths = [
            Path("../AAA").resolve(),
            Path("AAA").resolve(),
            AAA_ROOT,
            Path.cwd() / "AAA",
            Path.cwd().parent / "AAA",
        ]
        
        logger.info(f"🔍 当前工作目录: {Path.cwd()}")
        
        full_path = None
        tried_paths = []
        for base in base_paths:
            test_path = base / clean_path
            tried_paths.append(str(test_path))
            logger.info(f"🔍 尝试路径: {test_path} - 存在: {test_path.exists()}")
            if test_path.exists():
                full_path = test_path
                break
        
        if full_path is None:
            error_msg = f"文件不存在: {clean_path}\n尝试过的路径:\n" + "\n".join(f"  - {p}" for p in tried_paths)
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        
        logger.info(f"📄 输入文件: {full_path}")
        
        # 确定输出路径（始终复制文件，不直接修改原文件）
        if output_path:
            # 清理输出路径（去除引号、空格、统一分隔符）
            clean_output = output_path.strip().strip('"').strip("'")
            clean_output = clean_output.replace("\\", "/")
            
            if clean_output.startswith('/AAA/'):
                clean_output = clean_output[5:]
            elif clean_output.startswith('AAA/'):
                clean_output = clean_output[4:]
            elif clean_output.startswith('/'):
                clean_output = clean_output[1:]
            
            # 使用与输入文件相同的基础路径
            output_full_path = full_path.parent / Path(clean_output).name
            # 或者如果指定了完整相对路径，使用AAA_ROOT
            if "/" in clean_output:
                for base in base_paths:
                    if base.exists():
                        output_full_path = base / clean_output
                        break
        else:
            # 如果没有指定输出路径，生成一个带 _cleaned 后缀的文件名
            # 确保不会覆盖原文件
            file_stem = full_path.stem
            file_suffix = full_path.suffix
            output_full_path = full_path.parent / f"{file_stem}_cleaned{file_suffix}"
            logger.info(f"📋 未指定输出路径，将生成清理后的文件: {output_full_path}")
        
        # 如果输出路径和输入路径相同，自动生成不同的文件名
        if output_full_path.resolve() == full_path.resolve():
            file_stem = full_path.stem
            file_suffix = full_path.suffix
            output_full_path = full_path.parent / f"{file_stem}_cleaned{file_suffix}"
            logger.info(f"📋 输出路径与原文件相同，自动生成新文件名: {output_full_path}")
        
        # 先复制原文件到输出路径，确保原文件不被修改
        logger.info(f"📋 复制原文件到: {output_full_path}")
        output_full_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(full_path, output_full_path)
        logger.info(f"✅ 文件复制完成，原文件保持不变")
        
        logger.info(f"📄 输出文件: {output_full_path}")
        logger.info(f"🔧 删除首行: {remove_first_line}")
        logger.info(f"🔧 清理控件: {remove_content_controls}")
        
        # 执行清理（优先 python-docx，失败回退 COM）
        if remove_content_controls:
            # 先尝试 python-docx（快速）
            logger.info("🚀 尝试使用 python-docx 清理...")
            result = _clean_document_with_docx(
                file_path=str(output_full_path),
                output_path=str(output_full_path),
                remove_first_line=remove_first_line
            )

            # 如果失败，回退到 COM
            if not result.get("success"):
                logger.warning("⚠️ python-docx 清理失败，回退到 COM 方案")
                result = _clean_content_controls_preserve_content(
                    file_path=str(output_full_path),
                    output_path=str(output_full_path),
                    remove_first_line=remove_first_line
                )
        else:
            # 只删除首行，不清理控件
            result = {"success": False, "error": "仅删除首行功能暂未单独实现"}
            if remove_first_line:
                # 使用现有的清理首行函数
                with open(output_full_path, "rb") as f:
                    src_bytes = f.read()
                suffix = output_full_path.suffix
                cleaned_bytes = _clear_first_line_with_com_bytes(src_bytes, suffix)
                if cleaned_bytes:
                    with open(output_full_path, "wb") as f:
                        f.write(cleaned_bytes)
                    result = {
                        "success": True,
                        "controls_removed": 0,
                        "first_line_removed": True,
                        "output_file": str(output_full_path)
                    }
        
        if result.get("success"):
            logger.info("✅ 文档清理成功")
            logger.info(f"   - 清理控件: {result.get('controls_removed', 0)} 个")
            logger.info(f"   - 删除首行: {result.get('first_line_removed', False)}")
            
            # 构建返回的相对路径（相对于AAA）
            output_rel_path = None
            try:
                for base in [Path("../AAA"), Path("AAA"), AAA_ROOT]:
                    try:
                        output_rel_path = "AAA/" + str(Path(output_full_path).relative_to(base)).replace("\\", "/")
                        break
                    except ValueError:
                        continue
            except Exception:
                output_rel_path = str(output_full_path)
            
            return JSONResponse({
                "success": True,
                "output_file": output_rel_path or str(output_full_path),
                "controls_removed": result.get("controls_removed", 0),
                "first_line_removed": result.get("first_line_removed", False),
                "markers_removed": result.get("markers_removed", False)
            })
        else:
            raise HTTPException(status_code=500, detail=f"清理失败: {result.get('error', '未知错误')}")
    
    except FileNotFoundError as e:
        logger.error(f"文件不存在: {e}")
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"处理失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ========== 以下是原bundle接口的剩余代码（已注释） ==========
# 原接口代码已移至 app_bundle_backup.py 文件
"""
原bundle接口的主要流程：
1. 接收zip文件和JSON数据
2. 解压到临时目录  
3. 从解压文件中读取资源
4. 执行插入操作
5. 返回结果文档

现在直接模式的优势：
- 无需打包zip文件
- 直接使用共享文件夹AAA中的资源
- 更快速、更简单
"""


# ========== Linux转发的预处理接口（新增）==========

@app.post("/api/v1/preprocessing/process")
def preprocessing_process(
    request: Request,
    file_path: str = Form(..., description="文件相对路径（相对于AAA/project_data）"),
    folder_path: str = Form(..., description="项目文件夹路径"),
    filename: str = Form(..., description="文件名"),
    file_id: str = Form(None, description="文件ID"),
    force_ocr: bool = Form(False, description="是否强制OCR"),
    extract_regions: bool = Form(True, description="是否提取表格图片"),
    extract_assets: bool = Form(True, description="是否提取资产"),
    chunking_enabled: bool = Form(True, description="是否启用分块"),
    chunking_mode: str = Form("heading", description="分块模式"),
):
    """
    Linux转发的预处理接口（同步执行，避免阻塞事件循环）
    处理Word/RTF/Excel等需要Windows环境的文件
    从共享目录AAA读取文件
    """
    if not _auth_ok(request):
        raise HTTPException(status_code=401, detail="unauthorized")

    with _windows_serial_task_guard("preprocessing_process", request):
        return _preprocessing_process_impl(
            request=request,
            file_path=file_path,
            folder_path=folder_path,
            filename=filename,
            file_id=file_id,
            force_ocr=force_ocr,
            extract_regions=extract_regions,
            extract_assets=extract_assets,
            chunking_enabled=chunking_enabled,
            chunking_mode=chunking_mode,
        )


def _preprocessing_process_impl(
    *,
    request: Request,
    file_path: str,
    folder_path: str,
    filename: str,
    file_id: str | None,
    force_ocr: bool,
    extract_regions: bool,
    extract_assets: bool,
    chunking_enabled: bool,
    chunking_mode: str,
):
    try:
        logger.info(f"📥 收到预处理请求: {filename}")
        logger.info(f"   文件路径: {file_path}")
        
        # ✅ 延迟导入PreprocessingService（恢复原来的方式）
        from service.windows.preprocessing.service import PreprocessingService

        # ✅ 从共享目录读取文件
        AAA_ROOT = Path(os.getenv("WINDOWS_AAA_ROOT", "AAA"))
        full_file_path = AAA_ROOT / "project_data" / file_path
        
        if not full_file_path.exists():
            raise FileNotFoundError(f"文件不存在: {full_file_path}")
        
        logger.info(f"📁 找到文件: {full_file_path}")
        
        # 构建输出路径（共享目录）
        output_dir = AAA_ROOT / "Preprocessing" / folder_path / Path(filename).stem
        
        logger.info(f"🔄 开始预处理: {filename} → {output_dir}")
    
        # 调用预处理服务
        preprocessing_svc = PreprocessingService()
        extra_info = {"file_id": file_id} if file_id else {}
        
        result = preprocessing_svc.preprocess(
            file_path=full_file_path,
            force_ocr=force_ocr,
            extract_regions=extract_regions,
            extract_assets=extract_assets,
            chunking_enabled=chunking_enabled,
            chunking_mode=chunking_mode,
            output_dir=output_dir,
            extra_info=extra_info
        )
        
        # 构建返回路径（相对于AAA的路径，供Linux访问）
        preprocessed_file = str(result.work_dir / "preprocessed.json") if result.work_dir else None
        chunks_file = result.processing_info.get('structured_chunks_file', None)
        preprocessed_dir = str(result.work_dir) if result.work_dir else None
        
        # 转换为Linux可访问的AAA相对路径
        if preprocessed_file:
            p = Path(preprocessed_file)
            try:
                rel = p.relative_to(AAA_ROOT)
                preprocessed_file = "AAA/" + rel.as_posix()
            except Exception:
                preprocessed_file = p.as_posix()
        if chunks_file:
            p2 = Path(chunks_file)
            try:
                rel2 = p2.relative_to(AAA_ROOT)
                chunks_file = "AAA/" + rel2.as_posix()
            except Exception:
                chunks_file = p2.as_posix()
        if preprocessed_dir:
            pd = Path(preprocessed_dir)
            try:
                rel3 = pd.relative_to(AAA_ROOT)
                preprocessed_dir = "AAA/" + rel3.as_posix()
            except Exception:
                preprocessed_dir = pd.as_posix()
        
        logger.info(f"✅ 预处理成功: {filename}")
        
        return JSONResponse({
            "success": True,
            "id": file_id or "",
            "status": "success",
            "file_name": filename,
            "file_type": result.file_type.value if hasattr(result.file_type, 'value') else str(result.file_type),
            "preprocessed_json": preprocessed_file,
            "preprocessed_dir": preprocessed_dir,
            "chunks_file": chunks_file,
            "regions_count": len(result.regions) if result.regions else 0,
            "processing_method": "windows_server"
        })
        
    except Exception as e:
        logger.error(f"❌ 预处理失败: {e}", exc_info=True)
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "id": file_id or "",
                "status": "fail",
                "file_name": filename,
                "error_message": str(e)
            }
        )


if __name__ == "__main__":
    import uvicorn

    
    host = os.getenv("HOST", "0.0.0.0")
    port = int(os.getenv("PORT", "8081"))
    
    logger.info("=" * 70)
    logger.info("Windows Bridge Service 启动")
    logger.info(f"监听地址: {host}:{port}")
    logger.info("=" * 70)
    
    # 直接传 app 对象，避免 import path 解析失败
    uvicorn.run(app, host=host, port=port, reload=False, log_level="info")


