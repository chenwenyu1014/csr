#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word文档处理服务
跨平台的Word文档操作，支持占位符替换和内容插入
"""

import logging
import re
from pathlib import Path
import os
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx未安装，Word文档功能将不可用")

try:
    import win32com.client
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False

try:
    from service.linux.bridge.windows_bridge_client import WindowsBridgeClient  # type: ignore
except Exception:
    WindowsBridgeClient = None  # type: ignore

@dataclass
class WordInsertionResult:
    """Word插入结果"""
    success: bool
    message: str
    output_file: Optional[str] = None
    error: Optional[str] = None
    inserted_paragraphs: List[str] = None

class WordDocumentService:
    """Word文档处理服务"""
    
    def __init__(self, use_win32: bool = False):
        """
        初始化Word文档服务
        
        Args:
            use_win32: 是否使用Windows COM接口（仅Windows可用）
        """
        self.use_win32 = use_win32 and WIN32_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        
        if not self.docx_available and not self.use_win32:
            raise ImportError("需要安装python-docx库或Windows COM接口")
        
        logger.info(f"Word文档服务初始化完成 - 使用方式: {'Win32 COM' if self.use_win32 else 'python-docx'}")
    
    def insert_content_to_word(self,
                              template_file: str,
                              content_data: Dict[str, str],
                              output_file: Optional[str] = None,
                              placeholder_format: str = "{{%s}}") -> WordInsertionResult:
        """
        将生成的内容插入到Word文档的占位符中
        
        Args:
            template_file: Word模板文件路径
            content_data: 内容数据字典，key为段落ID，value为生成的内容
            output_file: 输出文件路径，如果为None则自动生成
            placeholder_format: 占位符格式，默认为"{{段落ID}}"
            
        Returns:
            WordInsertionResult: 插入结果
        """
        try:
            if self.use_win32:
                return self._insert_with_win32(template_file, content_data, output_file, placeholder_format)
            else:
                return self._insert_with_docx(template_file, content_data, output_file, placeholder_format)
                
        except Exception as e:
            logger.error(f"Word文档插入失败: {e}", exc_info=True)
            return WordInsertionResult(
                success=False,
                message="Word文档插入失败",
                error=str(e)
            )
    
    def _insert_with_docx(self,
                         template_file: str,
                         content_data: Dict[str, str],
                         output_file: Optional[str],
                         placeholder_format: str) -> WordInsertionResult:
        """使用python-docx库插入内容"""
        try:
            # 加载模板文档
            doc = Document(template_file)
            inserted_paragraphs = []
            
            def _replace_placeholder_preserve_style(paragraph, placeholder: str, content: str) -> bool:
                """在保持占位符样式（字体/字号/粗斜体等）的前提下进行替换。
                处理优先级：
                1) 占位符独占整段（最常见模板场景）：用首个 run 的样式替换整段文本；
                2) 占位符落在单个 run 内：仅替换该 run 的文本，样式天然保留；
                3) 兜底：段落文本级替换（可能丢失细粒度样式）。
                """
                try:
                    full_text = paragraph.text or ""
                    # 情况1：整段即占位符
                    if full_text.strip() == placeholder:
                        runs = list(getattr(paragraph, 'runs', []) or [])
                        if runs:
                            base = runs[0]
                            base.text = content
                            # 清空其余 run 文本，避免残留
                            for r in runs[1:]:
                                r.text = ""
                        else:
                            paragraph.add_run(content)
                        return True
                    # 情况2：占位符位于某个 run 内
                    for r in getattr(paragraph, 'runs', []) or []:
                        txt = r.text or ""
                        if placeholder in txt:
                            r.text = txt.replace(placeholder, content)
                            return True
                    # 情况3：兜底（不保证样式细节完全保留）
                    if placeholder in full_text:
                        paragraph.text = full_text.replace(placeholder, content)
                        return True
                except Exception:
                    # 任意异常回退到不保留样式的替换以保证可用性
                    ft = paragraph.text or ""
                    if placeholder in ft:
                        paragraph.text = ft.replace(placeholder, content)
                        return True
                return False

            # 遍历所有段落，查找占位符
            for paragraph in doc.paragraphs:
                for paragraph_id, content in content_data.items():
                    placeholder = placeholder_format % paragraph_id
                    if placeholder in (paragraph.text or ""):
                        if _replace_placeholder_preserve_style(paragraph, placeholder, content):
                            inserted_paragraphs.append(paragraph_id)
                            logger.info(f"已插入段落: {paragraph_id}")
            
            # 遍历所有表格，查找占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for paragraph_id, content in content_data.items():
                                placeholder = placeholder_format % paragraph_id
                                if placeholder in (paragraph.text or ""):
                                    if _replace_placeholder_preserve_style(paragraph, placeholder, content):
                                        inserted_paragraphs.append(paragraph_id)
                                        logger.info(f"已在表格中插入段落: {paragraph_id}")

            # 模板阶段不再插入分节符
            
            # 生成输出文件名
            if output_file is None:
                template_path = Path(template_file)
                output_file = str(template_path.parent / f"{template_path.stem}_filled{template_path.suffix}")
            
            # 保存文档
            doc.save(output_file)
            
            return WordInsertionResult(
                success=True,
                message=f"成功插入{len(inserted_paragraphs)}个段落到Word文档",
                output_file=output_file,
                inserted_paragraphs=inserted_paragraphs
            )
            
        except Exception as e:
            logger.error(f"python-docx插入失败: {e}", exc_info=True)
            return WordInsertionResult(
                success=False,
                message="python-docx插入失败",
                error=str(e)
            )
    
    def _insert_with_win32(self,
                          template_file: str,
                          content_data: Dict[str, str],
                          output_file: Optional[str],
                          placeholder_format: str) -> WordInsertionResult:
        """使用Windows COM接口插入内容"""
        try:
            # 启动Word应用程序
            from utils.windows_com import safe_dispatch
            word_app = safe_dispatch("Word.Application", use_ex=False, logger=logger)
            word_app.Visible = False
            
            # 打开模板文档
            doc = word_app.Documents.Open(template_file)
            inserted_paragraphs = []
            
            # 执行查找和替换
            for paragraph_id, content in content_data.items():
                placeholder = placeholder_format % paragraph_id
                
                # 查找占位符
                find_obj = doc.Content.Find
                find_obj.ClearFormatting()
                find_obj.Replacement.ClearFormatting()
                find_obj.Text = placeholder
                find_obj.Replacement.Text = content
                find_obj.Forward = True
                find_obj.Wrap = 1  # wdFindContinue
                find_obj.Format = False
                find_obj.MatchCase = False
                find_obj.MatchWholeWord = False
                find_obj.MatchWildcards = False
                find_obj.MatchSoundsLike = False
                find_obj.MatchAllWordForms = False
                
                # 执行替换
                result = find_obj.Execute(Replace=2)  # wdReplaceAll
                if result:
                    inserted_paragraphs.append(paragraph_id)
                    logger.info(f"已插入段落: {paragraph_id}")
            
            # 生成输出文件名
            if output_file is None:
                template_path = Path(template_file)
                output_file = str(template_path.parent / f"{template_path.stem}_filled{template_path.suffix}")
            
            # 保存文档
            doc.SaveAs(output_file)
            doc.Close()
            word_app.Quit()
            
            return WordInsertionResult(
                success=True,
                message=f"成功插入{len(inserted_paragraphs)}个段落到Word文档",
                output_file=output_file,
                inserted_paragraphs=inserted_paragraphs
            )
            
        except Exception as e:
            logger.error(f"Win32 COM插入失败: {e}", exc_info=True)
            try:
                if 'doc' in locals():
                    doc.Close()
                if 'word_app' in locals():
                    word_app.Quit()
            except:
                pass
            return WordInsertionResult(
                success=False,
                message="Win32 COM插入失败",
                error=str(e)
            )
    
    def create_template_with_placeholders(self,
                                        template_file: str,
                                        paragraph_ids: List[str],
                                        placeholder_format: str = "{{%s}}") -> bool:
        """
        创建包含占位符的Word模板文件
        
        Args:
            template_file: 模板文件路径
            paragraph_ids: 段落ID列表
            placeholder_format: 占位符格式
            
        Returns:
            bool: 是否创建成功
        """
        try:
            if not self.docx_available:
                logger.error("python-docx不可用，无法创建模板")
                return False
            
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('CSR文档模板', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加说明
            doc.add_paragraph("本模板包含以下占位符，系统将自动替换为生成的内容：")
            
            # 为每个段落ID添加占位符
            for paragraph_id in paragraph_ids:
                doc.add_heading(f'段落: {paragraph_id}', level=1)
                placeholder = placeholder_format % paragraph_id
                doc.add_paragraph(placeholder)
                doc.add_paragraph()  # 空行
            
            # 保存模板
            doc.save(template_file)
            logger.info(f"模板文件已创建: {template_file}")
            return True
            
        except Exception as e:
            logger.error(f"创建模板失败: {e}", exc_info=True)
            return False
    
    def validate_template(self, template_file: str, paragraph_ids: List[str], 
                         placeholder_format: str = "{{%s}}") -> Dict[str, Any]:
        """
        验证模板文件中的占位符
        
        Args:
            template_file: 模板文件路径
            paragraph_ids: 段落ID列表
            placeholder_format: 占位符格式
            
        Returns:
            Dict: 验证结果
        """
        try:
            if not self.docx_available:
                return {"valid": False, "error": "python-docx不可用"}
            
            doc = Document(template_file)
            found_placeholders = []
            missing_placeholders = []
            
            # 收集所有文本内容
            all_text = ""
            for paragraph in doc.paragraphs:
                all_text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            all_text += paragraph.text + "\n"
            
            # 检查每个段落ID的占位符
            for paragraph_id in paragraph_ids:
                placeholder = placeholder_format % paragraph_id
                if placeholder in all_text:
                    found_placeholders.append(paragraph_id)
                else:
                    missing_placeholders.append(paragraph_id)
            
            return {
                "valid": len(missing_placeholders) == 0,
                "found_placeholders": found_placeholders,
                "missing_placeholders": missing_placeholders,
                "total_expected": len(paragraph_ids),
                "total_found": len(found_placeholders)
            }
            
        except Exception as e:
            logger.error(f"验证模板失败: {e}", exc_info=True)
            return {"valid": False, "error": str(e)}
    
    def get_available_methods(self) -> List[str]:
        """获取可用的处理方法"""
        methods = []
        if self.docx_available:
            methods.append("python-docx")
        if WIN32_AVAILABLE:
            methods.append("win32-com")
        return methods
    
    def get_system_info(self) -> Dict[str, Any]:
        """获取系统信息"""
        return {
            "docx_available": self.docx_available,
            "win32_available": WIN32_AVAILABLE,
            "current_method": "win32-com" if self.use_win32 else "python-docx",
            "available_methods": self.get_available_methods()
        }

    def _create_short_path_for_com(self, original_path: Path, must_exist: bool = True) -> Path:
        """为COM操作创建短路径副本，解决长路径/中文路径问题。
        
        Args:
            original_path: 原始文件路径
            must_exist: 是否要求文件必须存在（False时仅生成临时路径，不复制文件）
        """
        try:
            s = str(original_path)
            # Python 3.8+ 字符串支持 isascii；为兼容性使用 getattr
            if len(s) < 250 and getattr(s, "isascii", lambda: False)():
                return original_path
        except Exception:
            pass

        # 将文件复制到系统临时目录下的 com_temp，并使用ASCII安全文件名
        import tempfile as _tempfile
        import shutil as _shutil
        import uuid as _uuid
        from pathlib import Path as _Path

        temp_dir = _Path(_tempfile.gettempdir()) / "com_temp"
        try:
            temp_dir.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass

        short_name = f"{original_path.stem[:10]}_{_uuid.uuid4().hex[:8]}{original_path.suffix}"
        temp_path = temp_dir / short_name
        
        if must_exist:
            try:
                _shutil.copy2(str(original_path), str(temp_path))
                logger.info(f"为COM操作创建临时短路径副本: {original_path} -> {temp_path}")
                return temp_path
            except Exception:
                # 复制失败则退回原始路径
                return original_path
        else:
            # 仅返回临时路径，不复制文件
            logger.info(f"为COM操作生成临时短路径: {temp_path}")
            return temp_path

    def scan_all_tag_regions(self, marked_file: str) -> List[Dict[str, str]]:
        """扫描标记版文档中的所有标签区间，返回可供导出的 regions 列表。

        优先策略：
        - 使用 Win32 COM 读取对象计数，以构造标准化的标签名（无需解析XML）。
        - 标签格式为：{{Table_i_Start}}/{{Table_i_End}} 与 {{Image_i_Start}}/{{Image_i_End}}
        - 注意：若文档同时存在 InlineShapes 与 Shapes，Image_i 可能存在重号；此处保持原始编号。
        """
        regions: List[Dict[str, str]] = []
        # 优先 WindowsBridge（默认开；可通过 DISABLE_WINDOWS_BRIDGE 或 USE_WINDOWS_BRIDGE/USE_WINDOWS_BRIDGE_WORD 扫描关闭）：
        try:
            if WindowsBridgeClient is not None:
                def _env_true(v: str) -> bool:
                    return (str(v).strip().lower() in ('1','true','yes','y','on','是'))
                disabled = _env_true(os.getenv('DISABLE_WINDOWS_BRIDGE', '0'))
                use_bridge_global = _env_true(os.getenv('USE_WINDOWS_BRIDGE', '1'))
                use_bridge_feature = _env_true(os.getenv('USE_WINDOWS_BRIDGE_WORD', '1'))
                if (not disabled) and use_bridge_global and use_bridge_feature:
                    cli = WindowsBridgeClient()
                    if cli.is_configured():
                        import json as _json
                        import requests as _requests
                        files = {"file": (Path(marked_file).name, open(marked_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                        url = f"{cli.base_url}/api/v1/word/scan_regions"
                        headers = cli._headers()  # type: ignore
                        resp = _requests.post(url, files=files, timeout=cli.timeout, headers=headers)
                        if resp.status_code == 200:
                            data = resp.json()
                            if isinstance(data, list):
                                return data  # 直接返回远程扫描结果
        except Exception:
            pass
        if WIN32_AVAILABLE:
            try:
                from pathlib import Path as _Path
                from utils.windows_com import safe_dispatch
                word_app = safe_dispatch("Word.Application", use_ex=False, logger=logger)
                word_app.Visible = False
                try:
                    word_app.DisplayAlerts = 0
                except Exception:
                    pass
                src_p = _Path(marked_file)
                try:
                    if not src_p.is_absolute():
                        src_p = src_p.resolve()
                except Exception:
                    pass
                com_src = self._create_short_path_for_com(src_p)
                doc = word_app.Documents.Open(str(com_src), ReadOnly=True)
                try:
                    try:
                        tcnt = int(doc.Tables.Count)
                    except Exception:
                        tcnt = 0
                    for i in range(1, tcnt + 1):
                        s = f"{{{{Table_{i}_Start}}}}"; e = f"{{{{Table_{i}_End}}}}"
                        regions.append({"start": s, "end": e, "name": s.replace("{", "").replace("}", "")})

                    try:
                        icnt = int(doc.InlineShapes.Count)
                    except Exception:
                        icnt = 0
                    for i in range(1, icnt + 1):
                        s = f"{{{{Image_{i}_Start}}}}"; e = f"{{{{Image_{i}_End}}}}"
                        regions.append({"start": s, "end": e, "name": s.replace("{", "").replace("}", "")})

                    try:
                        scnt = int(doc.Shapes.Count)
                    except Exception:
                        scnt = 0
                    for i in range(1, scnt + 1):
                        s = f"{{{{Image_{i}_Start}}}}"; e = f"{{{{Image_{i}_End}}}}"
                        regions.append({"start": s, "end": e, "name": s.replace("{", "").replace("}", "")})
                finally:
                    try:
                        doc.Close(SaveChanges=False)
                    except Exception:
                        pass
                    try:
                        word_app.Quit()
                    except Exception:
                        pass
                    try:
                        if com_src != src_p and Path(str(com_src)).exists():
                            Path(str(com_src)).unlink()
                    except Exception:
                        pass
                return regions
            except Exception:
                pass

        # 回退：若无法使用 COM，则尝试基于 python-docx 解析段落文本（可能遗漏表格内标签）
        try:
            if not self.docx_available:
                return regions
            import re as _re
            from docx import Document as _Document  # type: ignore
            doc = _Document(marked_file)
            starts: List[tuple[int, str]] = []
            ends: Dict[str, int] = {}
            pat_s = _re.compile(r"^\{\{(Table|Image)_(\d+)_Start\}\}$")
            pat_e = _re.compile(r"^\{\{(Table|Image)_(\d+)_End\}\}$")
            for idx, p in enumerate(doc.paragraphs):
                txt = (p.text or "").strip()
                m = pat_s.match(txt)
                if m:
                    key = f"{m.group(1)}_{m.group(2)}"
                    starts.append((idx, key))
                m2 = pat_e.match(txt)
                if m2:
                    key = f"{m2.group(1)}_{m2.group(2)}"
                    ends[key] = idx
            for s_idx, key in starts:
                e_idx = ends.get(key)
                if e_idx is not None and e_idx > s_idx:
                    s = f"{{{{{key}_Start}}}}"; e = f"{{{{{key}_End}}}}"
                    regions.append({"start": s, "end": e, "name": s.replace("{", "").replace("}", "")})
        except Exception:
            pass
        return regions

    def mark_tables_and_images_with_com(self, source_file: str, marked_file: Optional[str] = None) -> Optional[str]:
        """
        使用Word原生COM接口在文档中插入表/图起止标签，并保存为 *_marked.docx。
        仅插入标签，不改变内容；表用 Tables，内嵌图用 InlineShapes，浮动图用 Shapes.Anchor 段落。

        Args:
            source_file: 原始Word路径
            marked_file: 目标标记文件路径（可选）

        Returns:
            str|None: 成功返回标记文件完整路径
        """
        # 优先 WindowsBridge 远程执行（默认开；可通过 DISABLE_WINDOWS_BRIDGE 或 USE_WINDOWS_BRIDGE/USE_WINDOWS_BRIDGE_WORD 关闭）
        try:
            if WindowsBridgeClient is not None:
                def _env_true(v: str) -> bool:
                    return (str(v).strip().lower() in ('1','true','yes','y','on','是'))
                disabled = _env_true(os.getenv('DISABLE_WINDOWS_BRIDGE', '0'))
                use_bridge_global = _env_true(os.getenv('USE_WINDOWS_BRIDGE', '1'))
                use_bridge_feature = _env_true(os.getenv('USE_WINDOWS_BRIDGE_WORD', '1'))
                if (not disabled) and use_bridge_global and use_bridge_feature:
                    cli = WindowsBridgeClient()
                    if cli.is_configured():
                        import requests as _requests
                        from pathlib import Path as _Path
                        src_p = _Path(source_file)
                        files = {"file": (src_p.name, open(source_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                        url = f"{cli.base_url}/api/v1/word/mark_tables_images"
                        headers = cli._headers()  # type: ignore
                        resp = _requests.post(url, files=files, timeout=cli.timeout, headers=headers)
                        if resp.status_code == 200 and resp.content:
                            if not marked_file:
                                marked_p = src_p.parent / f"{src_p.stem}_marked{src_p.suffix}"
                            else:
                                marked_p = _Path(marked_file)
                            marked_p.write_bytes(resp.content)
                            return str(marked_p)
        except Exception:
            pass

        try:
            if not WIN32_AVAILABLE:
                return None
            from pathlib import Path as _Path
            src_p = _Path(source_file)
            if not marked_file:
                marked_p = src_p.parent / f"{src_p.stem}_marked{src_p.suffix}"
            else:
                marked_p = _Path(marked_file)

            from utils.windows_com import safe_dispatch
            word_app = safe_dispatch("Word.Application", use_ex=False, logger=logger)
            word_app.Visible = False
            word_app.ScreenUpdating = False
            # 为COM操作准备安全路径
            _safe_src = self._create_short_path_for_com(src_p, must_exist=True)
            _safe_marked = self._create_short_path_for_com(marked_p, must_exist=False) if len(str(marked_p)) > 200 else marked_p
            doc = word_app.Documents.Open(str(_safe_src))
            # 保存为副本并在副本上标记
            doc.SaveAs(str(_safe_marked))

            # 表格：按出现顺序从后往前插入标签，避免位置偏移影响后续索引
            try:
                tcnt = int(doc.Tables.Count)
            except Exception:
                tcnt = 0
            for i in range(tcnt, 0, -1):
                try:
                    tbl = doc.Tables.Item(i)
                    try:
                        start_tag = f"{{{{Table_{i}_Start}}}}"
                        end_tag = f"{{{{Table_{i}_End}}}}"
                        # before
                        rs = doc.Range(Start=tbl.Range.Start, End=tbl.Range.Start)
                        rs.InsertAfter(start_tag)
                        rs.InsertParagraphAfter()
                        # after
                        re = doc.Range(Start=tbl.Range.End, End=tbl.Range.End)
                        re.InsertParagraphBefore()
                        re.InsertBefore(end_tag)
                    except Exception:
                        pass
                except Exception:
                    pass

            # 内嵌图片：InlineShapes
            try:
                icnt = int(doc.InlineShapes.Count)
            except Exception:
                icnt = 0
            for i in range(icnt, 0, -1):
                try:
                    ish = doc.InlineShapes.Item(i)
                    start_tag = f"{{{{Image_{i}_Start}}}}"
                    end_tag = f"{{{{Image_{i}_End}}}}"
                    rs = doc.Range(Start=ish.Range.Start, End=ish.Range.Start)
                    rs.InsertAfter(start_tag)
                    rs.InsertParagraphAfter()
                    re = doc.Range(Start=ish.Range.End, End=ish.Range.End)
                    re.InsertParagraphBefore()
                    re.InsertBefore(end_tag)
                except Exception:
                    pass

            # 浮动图片：Shapes（按锚点段落插入标签）
            try:
                scnt = int(doc.Shapes.Count)
            except Exception:
                scnt = 0
            for i in range(scnt, 0, -1):
                try:
                    sh = doc.Shapes.Item(i)
                    start_tag = f"{{{{Image_{i}_Start}}}}"
                    end_tag = f"{{{{Image_{i}_End}}}}"
                    anc = sh.Anchor
                    rstart = doc.Range(Start=anc.Start, End=anc.Start)
                    rstart.InsertAfter(start_tag)
                    rstart.InsertParagraphAfter()
                    rend = doc.Range(Start=anc.End, End=anc.End)
                    rend.InsertParagraphBefore()
                    rend.InsertBefore(end_tag)
                except Exception:
                    pass

            doc.Save()
            doc.Close(SaveChanges=True)
            word_app.Quit()
            
            # 如果使用了临时短路径，需要复制回原始路径
            import shutil as _shutil
            try:
                if _safe_marked != marked_p:
                    _shutil.copy2(str(_safe_marked), str(marked_p))
                    Path(str(_safe_marked)).unlink()
            except Exception as e:
                logger.warning(f"复制标记文件失败: {e}")
            
            # 清理临时短路径
            try:
                if _safe_src != src_p and Path(str(_safe_src)).exists():
                    Path(str(_safe_src)).unlink()
            except Exception:
                pass
            return str(marked_p)
        except Exception as e:
            try:
                word_app.Quit()
            except Exception:
                pass
            logger.warning(f"COM标记失败: {e}")
            return None

    def export_regions_to_word(self,
                               marked_file: str,
                               regions: List[Dict[str, str]],
                               export_dir: str) -> List[Dict[str, str]]:
        """
        按给定的起止标签区间导出为多个独立的Word文件（优先使用Win32 COM，保证表/图完整复制）。

        Args:
            marked_file: 已插入 {{Table_X_Start}}/{{Table_X_End}} 与 {{Image_X_Start}}/{{Image_X_End}} 的标记版Word路径
            regions: 区间列表，每项包含 {"start": 开始标签文本, "end": 结束标签文本, "name": 导出文件基名}
            export_dir: 导出目录

        Returns:
            List[Dict[str, str]]: 每个导出的文件信息 {"name": name, "path": full_path}
        """
        results: List[Dict[str, str]] = []
        # 控制是否进行“断链、重嵌、兜底”图片处理；按需开启
        ENABLE_EMBED_FIX = False
        try:
            from pathlib import Path as _Path
            _Path(export_dir).mkdir(parents=True, exist_ok=True)
        except Exception:
            pass

        # 优先尝试Win32 COM，能够完整保留图片/表格等复杂对象
        if WIN32_AVAILABLE:
            try:
                from utils.windows_com import safe_dispatch
                word_app = safe_dispatch("Word.Application", use_ex=False, logger=logger)
                word_app.Visible = False
                word_app.ScreenUpdating = False
                try:
                    word_app.DisplayAlerts = 0
                except Exception:
                    pass
                # 规范化源文件路径：创建系统临时目录下的ASCII短路径副本供 COM 打开
                from pathlib import Path as _Path
                _src_p = _Path(marked_file)
                try:
                    if not _src_p.is_absolute():
                        _src_p = _src_p.resolve()
                except Exception:
                    pass
                _com_src = self._create_short_path_for_com(_src_p)
                try:
                    _src_mtime = _src_p.stat().st_mtime
                except Exception:
                    _src_mtime = None

                # 源文档以只读方式打开，单文档内完成查找与区间复制，避免频繁克隆带来的断连
                doc = word_app.Documents.Open(str(_com_src), ReadOnly=True)
                for reg in regions:
                    try:
                        start_txt = reg.get("start", "")
                        end_txt = reg.get("end", "")
                        base_name = reg.get("name", "region")
                        # 预计算输出路径并检查缓存
                        safe_name = base_name.replace("\\", "_").replace("/", "_").replace(":", "_")\
                                               .replace("*", "_").replace("?", "_").replace("\"", "_")\
                                               .replace("<", "_").replace(">", "_").replace("|", "_")
                        out_path = str(_Path(export_dir) / f"{safe_name}.docx")
                        try:
                            if _src_mtime is not None and _Path(out_path).exists() and _Path(out_path).stat().st_mtime >= _src_mtime:
                                # 缓存命中，直接返回
                                results.append({"name": base_name, "path": out_path})
                                continue
                        except Exception:
                            pass
                        # 1) 在只读的源文档内定位起止标签（使用独立Range，避免相互影响）
                        rng_s = doc.Content.Duplicate
                        fs = rng_s.Find
                        fs.ClearFormatting(); fs.Replacement.ClearFormatting()
                        fs.Text = start_txt
                        if not fs.Execute():
                            continue
                        s_pos = rng_s.End  # 开始标签之后

                        rng_e_scope = doc.Range(Start=s_pos, End=doc.Content.End)
                        rng_e = rng_e_scope.Duplicate
                        fe = rng_e.Find
                        fe.ClearFormatting(); fe.Replacement.ClearFormatting()
                        fe.Text = end_txt
                        if not fe.Execute():
                            continue
                        e_pos = rng_e.Start  # 结束标签起始
                        if e_pos <= s_pos:
                            continue

                        # 2) 新建空白文档，以源区间 Range.FormattedText 直接覆盖全文（避免删除操作导致的COM不稳定）
                        new_doc = word_app.Documents.Add()
                        try:
                            seg = doc.Range(Start=s_pos, End=e_pos)
                            new_doc.Content.FormattedText = seg.FormattedText
                        except Exception:
                            new_doc.Close(SaveChanges=False)
                            continue

                        # 3) 保存（使用临时短路径）
                        import tempfile as _tempfile
                        import shutil as _shutil
                        _out_p = _Path(out_path)
                        if len(str(_out_p)) > 200 or not str(_out_p).isascii():
                            # 使用临时短路径保存
                            _temp_dir = _Path(_tempfile.gettempdir()) / "com_temp"
                            _temp_dir.mkdir(parents=True, exist_ok=True)
                            import uuid as _uuid
                            _temp_out = _temp_dir / f"temp_{_uuid.uuid4().hex[:8]}.docx"
                            new_doc.SaveAs(str(_temp_out))
                            new_doc.Close(SaveChanges=True)
                            # 复制到目标路径
                            try:
                                _shutil.copy2(str(_temp_out), str(_out_p))
                                _temp_out.unlink()
                            except Exception as copy_err:
                                logger.warning(f"复制导出文件失败: {copy_err}")
                                continue
                        else:
                            # 直接保存
                            new_doc.SaveAs(out_path)
                            new_doc.Close(SaveChanges=True)
                        results.append({"name": base_name, "path": out_path})
                    except Exception as _e:
                        try:
                            if 'new_doc' in locals():
                                new_doc.Close(SaveChanges=False)
                        except Exception:
                            pass
                        logger.warning(f"导出区间失败（COM）: {reg}: {_e}")
                # 关闭文档与应用
                try:
                    doc.Close(SaveChanges=False)
                except Exception:
                    pass
                try:
                    word_app.Quit()
                except Exception:
                    pass
                # 清理临时短路径
                try:
                    if _com_src != _src_p and Path(str(_com_src)).exists():
                        Path(str(_com_src)).unlink()
                except Exception:
                    pass
                return results
            except Exception as e:
                logger.warning(f"使用Win32 COM导出失败，尝试python-docx回退: {e}")

        # 回退：python-docx 复制底层XML（图片可能因关系丢失无法保留，表格可正常复制）
        if not self.docx_available:
            logger.error("python-docx不可用，且COM失败，无法导出")
            return results

        try:
            from docx import Document as _Document
            import copy as _copy
            from pathlib import Path as _Path
            doc = _Document(marked_file)
            body = doc._element.body
            elements = list(body)
            try:
                _src_p = _Path(marked_file)
                _src_mtime = _src_p.stat().st_mtime
            except Exception:
                _src_mtime = None

            def _p_text(el) -> str:
                try:
                    if not str(getattr(el, 'tag', '')).endswith('p'):
                        return ""
                    texts: List[str] = []
                    for node in el.iter():
                        tag = str(getattr(node, 'tag', '') or '')
                        if tag.endswith('t') and getattr(node, 'text', None):
                            texts.append(node.text)
                    return ''.join(texts)
                except Exception:
                    return ""

            for reg in regions:
                try:
                    start_txt = reg.get("start", "")
                    end_txt = reg.get("end", "")
                    base_name = reg.get("name", "region")
                    # 预计算输出路径并检查缓存
                    safe_name = base_name.replace("\\", "_").replace("/", "_").replace(":", "_")\
                                           .replace("*", "_").replace("?", "_").replace("\"", "_")\
                                           .replace("<", "_").replace(">", "_").replace("|", "_")
                    out_path = str(_Path(export_dir) / f"{safe_name}.docx")
                    try:
                        if _src_mtime is not None and _Path(out_path).exists() and _Path(out_path).stat().st_mtime >= _src_mtime:
                            results.append({"name": base_name, "path": out_path})
                            continue
                    except Exception:
                        pass

                    start_idx = None
                    end_idx = None
                    for i, el in enumerate(elements):
                        if _p_text(el) == start_txt:
                            start_idx = i
                            break
                    if start_idx is None:
                        continue
                    for j in range(start_idx + 1, len(elements)):
                        if _p_text(elements[j]) == end_txt:
                            end_idx = j
                            break
                    if end_idx is None or end_idx <= start_idx + 1:
                        continue

                    seg = elements[start_idx + 1:end_idx]
                    new_doc = _Document()
                    # 清空默认正文
                    try:
                        new_doc._element.body.clear()
                    except Exception:
                        pass
                    for el in seg:
                        try:
                            new_doc._element.body.append(_copy.deepcopy(el))
                        except Exception:
                            pass

                    new_doc.save(out_path)
                    results.append({"name": base_name, "path": out_path})
                except Exception as _e:
                    logger.warning(f"导出区间失败（docx）: {reg}: {_e}")

        except Exception as e:
            logger.error(f"回退导出失败: {e}", exc_info=True)

        return results


    def export_all_tables_and_images(self, source_file: str, export_dir: str) -> List[Dict[str, str]]:
        """
        直接将文档内所有表格与图片各自导出为独立的 Word 文件（不依赖标签）。

        优先使用 Win32 COM：
        - 表格：使用 Range.FormattedText 复制，尽量保留格式
        - 内嵌图片与浮动图片：选择复制，粘贴到新文档

        回退到 python-docx：
        - 仅能导出表格（图片支持有限）
        """
        results: List[Dict[str, str]] = []
        try:
            from pathlib import Path as _Path
            _Path(export_dir).mkdir(parents=True, exist_ok=True)
        except Exception:
            pass

        def _sanitize(name: str) -> str:
            return (name.replace("\\", "_").replace("/", "_").replace(":", "_")
                        .replace("*", "_").replace("?", "_").replace("\"", "_")
                        .replace("<", "_").replace(">", "_").replace("|", "_"))

        # Win32 COM 路径（完全依赖原生：为每个对象克隆源文档并裁剪区间，保留主题与样式）
        if WIN32_AVAILABLE:
            try:
                from utils.windows_com import safe_dispatch
                word_app = safe_dispatch("Word.Application", use_ex=False, logger=logger)
                word_app.Visible = False
                try:
                    word_app.ScreenUpdating = False
                except Exception:
                    pass
                try:
                    word_app.DisplayAlerts = 0
                except Exception:
                    pass

                # 为包含中文/超长路径的文件创建 ASCII 短路径副本供 COM 打开
                from pathlib import Path as _Path
                _src_p = _Path(source_file)
                try:
                    if not _src_p.is_absolute():
                        _src_p = _src_p.resolve()
                except Exception:
                    pass
                _com_src = self._create_short_path_for_com(_src_p)

                # 先读取计数
                doc_probe = word_app.Documents.Open(str(_com_src), ReadOnly=True)
                try:
                    tcnt = int(doc_probe.Tables.Count)
                except Exception:
                    tcnt = 0
                try:
                    icnt = int(doc_probe.InlineShapes.Count)
                except Exception:
                    icnt = 0
                try:
                    scnt = int(doc_probe.Shapes.Count)
                except Exception:
                    scnt = 0
                try:
                    doc_probe.Close(SaveChanges=False)
                except Exception:
                    pass

                def _clear_headers_footers(_doc):
                    try:
                        wdHeaderFooterPrimary = 1
                        wdHeaderFooterFirstPage = 2
                        wdHeaderFooterEvenPages = 3
                        for si in range(1, int(_doc.Sections.Count) + 1):
                            try:
                                sec = _doc.Sections(si)
                            except Exception:
                                continue
                            # 头部
                            for hf in (wdHeaderFooterPrimary, wdHeaderFooterFirstPage, wdHeaderFooterEvenPages):
                                try:
                                    sec.Headers(hf).Range.Delete()
                                except Exception:
                                    pass
                            # 底部
                            for hf in (wdHeaderFooterPrimary, wdHeaderFooterFirstPage, wdHeaderFooterEvenPages):
                                try:
                                    sec.Footers(hf).Range.Delete()
                                except Exception:
                                    pass
                    except Exception:
                        pass

                # 表格：克隆->以表格 Range 覆盖全文（避免删除 Range 的权限问题）
                for i in range(1, tcnt + 1):
                    new_doc = None
                    try:
                        new_doc = word_app.Documents.Open(str(_com_src))
                        try:
                            tbl = new_doc.Tables.Item(i)
                        except Exception:
                            try:
                                # 若不可索引，跳过
                                new_doc.Close(SaveChanges=False)
                            except Exception:
                                pass
                            continue
                        keep_rng = new_doc.Range(Start=int(tbl.Range.Start), End=int(tbl.Range.End))
                        try:
                            new_doc.Content.FormattedText = keep_rng.FormattedText
                        except Exception:
                            # 回退到删除法
                            try:
                                s_pos = int(tbl.Range.Start); e_pos = int(tbl.Range.End)
                                new_doc.Range(Start=e_pos, End=new_doc.Content.End).Delete()
                                new_doc.Range(Start=new_doc.Content.Start, End=s_pos).Delete()
                            except Exception:
                                pass
                        _clear_headers_footers(new_doc)
                        out_path = str(_Path(export_dir) / f"{_sanitize(f'Table_{i}')}.docx")
                        new_doc.SaveAs(out_path)
                        new_doc.Close(SaveChanges=True)
                        results.append({"name": f"Table_{i}", "path": out_path})
                    except Exception:
                        try:
                            if new_doc is not None:
                                new_doc.Close(SaveChanges=False)
                        except Exception:
                            pass
                        continue

                # 内嵌图片：从只读源取 Range -> 新建空白文档以 Range 覆盖全文（避免编辑限制）
                for i in range(1, icnt + 1):
                    src_doc = None
                    new_doc = None
                    try:
                        src_doc = word_app.Documents.Open(str(_com_src), ReadOnly=True)
                        try:
                            ish = src_doc.InlineShapes.Item(i)
                        except Exception:
                            try:
                                src_doc.Close(SaveChanges=False)
                            except Exception:
                                pass
                            continue
                        keep_rng = src_doc.Range(Start=int(ish.Range.Start), End=int(ish.Range.End))
                        new_doc = word_app.Documents.Add()
                        try:
                            new_doc.Content.FormattedText = keep_rng.FormattedText
                        except Exception:
                            pass
                        _clear_headers_footers(new_doc)
                        out_path = str(_Path(export_dir) / f"{_sanitize(f'Image_{i}')}.docx")
                        new_doc.SaveAs(out_path)
                        new_doc.Close(SaveChanges=True)
                        results.append({"name": f"Image_{i}", "path": out_path})
                    except Exception:
                        try:
                            if new_doc is not None:
                                new_doc.Close(SaveChanges=False)
                        except Exception:
                            pass
                        try:
                            if src_doc is not None:
                                src_doc.Close(SaveChanges=False)
                        except Exception:
                            pass
                        continue

                # 浮动图片：从只读源取锚点段落 Range -> 新建空白文档以 Range 覆盖全文
                for i in range(1, scnt + 1):
                    src_doc = None
                    new_doc = None
                    try:
                        src_doc = word_app.Documents.Open(str(_com_src), ReadOnly=True)
                        try:
                            sh = src_doc.Shapes.Item(i)
                            anc = sh.Anchor
                        except Exception:
                            try:
                                src_doc.Close(SaveChanges=False)
                            except Exception:
                                pass
                            continue
                        keep_rng = src_doc.Range(Start=int(anc.Start), End=int(anc.End))
                        new_doc = word_app.Documents.Add()
                        try:
                            new_doc.Content.FormattedText = keep_rng.FormattedText
                        except Exception:
                            pass
                        _clear_headers_footers(new_doc)
                        out_path = str(_Path(export_dir) / f"{_sanitize(f'Image_{i}_F')}.docx")
                        new_doc.SaveAs(out_path)
                        new_doc.Close(SaveChanges=True)
                        results.append({"name": f"Image_{i}_F", "path": out_path})
                    except Exception:
                        try:
                            if new_doc is not None:
                                new_doc.Close(SaveChanges=False)
                        except Exception:
                            pass
                        try:
                            if src_doc is not None:
                                src_doc.Close(SaveChanges=False)
                        except Exception:
                            pass
                        continue

                try:
                    word_app.Quit()
                except Exception:
                    pass
                # 清理临时短路径
                try:
                    if _com_src != _src_p and Path(str(_com_src)).exists():
                        Path(str(_com_src)).unlink()
                except Exception:
                    pass
                return results
            except Exception as e:
                logger.warning(f"export_all_tables_and_images 使用COM失败: {e}")

        # 回退：python-docx（主要导出表格）
        if not self.docx_available:
            logger.error("python-docx不可用，且COM失败，无法导出全部图表")
            return results

        try:
            from docx import Document as _Document
            import copy as _copy
            doc = _Document(source_file)
            # 表格导出
            idx = 1
            for tbl in doc.tables:
                try:
                    new_doc = _Document()
                    # 清空
                    try:
                        new_doc._element.body.clear()
                    except Exception:
                        pass
                    # 复制表格底层XML
                    new_doc._element.body.append(_copy.deepcopy(tbl._element))
                    out_path = str((_Path(export_dir) / f"{_sanitize(f'Table_{idx}')}.docx"))
                    new_doc.save(out_path)
                    results.append({"name": f"Table_{idx}", "path": out_path})
                except Exception:
                    pass
                finally:
                    idx += 1
        except Exception as e:
            logger.error(f"export_all_tables_and_images 回退失败: {e}", exc_info=True)

        return results

# 全局服务实例
word_document_service = WordDocumentService()
