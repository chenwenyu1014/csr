#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文件预处理器 (File Processor)

统一入口：处理任意格式文件 → 标准化数据包

使用示例：
    processor = FileProcessor()
    result = processor.process("input.pdf")
    # 或
    result = processor.process("scanned.pdf", force_ocr=True)
    
    # 访问标准化数据
    print(result.text_content)
    print(result.regions)  # 拆分的片段
    print(result.assets)   # 提取的图片
"""

import logging
import os
import json
import threading
from pathlib import Path
from utils.output_manager import save_json, save_text
from typing import Optional, Dict, List, Any
from dataclasses import dataclass, field
from enum import Enum

logger = logging.getLogger(__name__)

# Word XML 命名空间
_WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_WNS_P = f'{{{_WNS}}}'

# LibreOffice 不支持多实例并发，所有 soffice 调用必须串行
_LIBREOFFICE_LOCK = threading.Lock()


class ContentType(Enum):
    """内容类型"""
    WORD = "word"           # 原生Word格式
    MARKDOWN = "markdown"   # Markdown格式
    STRUCTURED = "structured"  # 结构化JSON
    PLAIN_TEXT = "text"     # 纯文本


class FileType(Enum):
    """文件类型"""
    WORD_DOCX = "docx"
    WORD_DOC = "doc"
    RTF = "rtf"
    PDF = "pdf"
    EXCEL_XLSX = "xlsx"
    EXCEL_XLS = "xls"
    MARKDOWN = "markdown"
    UNKNOWN = "unknown"


@dataclass
class PreprocessedDocument:
    """预处理后的标准化文档"""
    
    # 元数据
    source_file: str
    file_type: FileType
    content_type: ContentType
    
    # 主要内容
    text_content: str = ""           # 纯文本内容
    markdown_content: str = ""       # Markdown格式内容
    structured_data: Dict = field(default_factory=dict)  # 结构化数据
    
    # 拆分的片段（按标签切分）
    regions: List[Dict[str, Any]] = field(default_factory=list)
    # [{
    #   "name": "Table_1",
    #   "type": "table",
    #   "content": "...",
    #   "file_path": "...",
    # }]
    
    # 提取的资源（图片、表格等）
    assets: List[Dict[str, Any]] = field(default_factory=list)
    # [{
    #   "type": "image",
    #   "name": "image_001.png",
    #   "path": "/path/to/image.png",
    #   "source_region": "Image_1"
    # }]
    
    # 处理信息
    processing_info: Dict[str, Any] = field(default_factory=dict)
    # {
    #   "conversion_method": "word_com",
    #   "ocr_used": False,
    #   "total_pages": 10,
    #   "processing_time": 2.5
    # }
    
    # 工作目录（临时文件存放）
    work_dir: Optional[Path] = None


class FileProcessor:
    """
    文件预处理器（统一入口）
    
    职责：
    1. 检测文件类型
    2. 选择最佳转换方案
    3. 执行格式转换
    4. 标记与切分
    5. 提取资源
    6. 返回标准化数据包
    """
    
    def __init__(self, cache_dir: str = "AAA/cache/preprocessing"):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def _stage_paths(self, work_dir: Path, stage: str) -> Dict[str, Path]:
        return {
            "done": work_dir / f"{stage}.done",
            "manifest": work_dir / f"{stage}_manifest.json",
        }

    def _ocr_strict(self) -> bool:
        """是否启用严格OCR模式（失败不回退）。通过环境变量 OCR_STRICT 控制。"""
        try:
            return str(os.getenv('OCR_STRICT', '0')).strip().lower() in ('1', 'true', 'yes', 'y', 'on')
        except Exception:
            return False

    def process(self, 
                file_path: str | Path,
                force_ocr: bool = False,
                extract_regions: bool = True,
                extract_assets: bool = True,
                content_type_override: Optional[str] = None,
                output_dir: Optional[str | Path] = None,
                extra_info: Optional[dict] = None) -> PreprocessedDocument:
        """
        处理文件并返回标准化数据包
        
        Args:
            file_path: 输入文件路径
            force_ocr: 强制使用OCR（即使是正常PDF）
            extract_regions: 是否拆分区域
            extract_assets: 是否提取资源（图片等）
            content_type_override: 手动指定的内容类型（覆盖自动分类）
                                  Word文档支持: 'txt'/'text'（文本）, 'jpg'/'images'（纯图）, 'list'/'tables'（纯表）
            output_dir: 输出目录（如果指定，直接在此目录处理；否则使用临时目录）
            extra_info: 额外信息字典，如 {"file_id": "xxx"}
            
        Returns:
            PreprocessedDocument: 标准化数据包
        """
        file_path = Path(file_path)
        logger.info(f"开始预处理: {file_path}")
        if content_type_override:
            logger.info(f"使用手动指定的内容类型: {content_type_override}")
        
        # ✅ 处理额外信息（file_id等）
        file_id = None
        if extra_info and isinstance(extra_info, dict):
            file_id = extra_info.get('file_id')
            if file_id:
                logger.info(f"文件ID: {file_id}")
        
        import time
        start_time = time.time()
        
        # 1. 检测文件类型
        file_type = self._detect_file_type(file_path, force_ocr)
        logger.info(f"文件类型: {file_type.value}")
        
        # 2. 创建工作目录
        if output_dir:
            # 使用指定的输出目录
            work_dir = Path(output_dir)
            work_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"使用指定输出目录: {work_dir}")
        else:
            # 使用临时目录
            import tempfile
            work_dir = Path(tempfile.mkdtemp(prefix=f"{file_path.stem}_"))
            logger.info(f"使用临时目录: {work_dir}")
        
        # 3. 转换为中间格式（统一为Markdown或纯文本）
        converted = self._convert_to_intermediate(file_path, file_type, work_dir, content_type_override)

        # 4. 获取Markdown内容；并在此处执行全局图片提取与替换，避免Base64导致正文过长
        marked_content = converted['content']
        global_md_images: List[Dict[str, Any]] = []
        try:
            if converted.get('content_type') == ContentType.MARKDOWN and marked_content:
                from service.windows.preprocessing.preprocessing_function.markdown.markdown_image_extractor import extract_base64_images, estimate_markdown_size_reduction
                cleaned_md, extracted = extract_base64_images(marked_content, output_dir=work_dir, image_subdir="images")
                if extracted:
                    marked_content = cleaned_md
                    # 记录统计信息，便于诊断
                    stats = estimate_markdown_size_reduction(converted['content'])
                    try:
                        converted.setdefault('metadata', {})['global_images_extracted'] = len(extracted)
                        converted['metadata']['images_dir'] = str((work_dir / 'images').resolve())
                        converted['metadata']['markdown_size_before'] = stats.get('original_size')
                        converted['metadata']['markdown_size_after_estimated'] = stats.get('estimated_cleaned_size')
                        converted['metadata']['markdown_size_reduction_percent'] = stats.get('reduction_percent')
                    except Exception:
                        pass
                    # 保存到本地变量，稍后加入assets
                    global_md_images = extracted
        except Exception:
            # 图片提取失败不影响主流程；保留原Markdown
            pass
        
        # 5. 获取区域（直接使用word_regions或excel_regions，不再从Markdown拆分）
        regions = []
        markdown_files = []  # 保存所有MD文件路径
        if extract_regions:
            # Excel：按sheet拆分
            try:
                excel_regions = converted.get('excel_regions')
                excel_md_regions = converted.get('excel_markdown_regions', [])
                if excel_regions:
                    # 为Excel regions构建正确的Label和path
                    formatted_regions = []
                    for i, region in enumerate(excel_regions):
                        sheet_name = region.get('sheet_name', f'Sheet{i+1}')
                        # 查找对应的markdown文件
                        md_path = None
                        for md_region in excel_md_regions:
                            if md_region.get('sheet_name') == sheet_name:
                                md_path = md_region.get('markdown_file')
                                if md_path and md_path not in markdown_files:
                                    markdown_files.append(md_path)
                                break
                        
                        formatted_regions.append({
                            'Label': f'Sheet_{sheet_name}',
                            'path': region.get('file_path') or '',  # Excel文件路径
                            'markdown_path': md_path or None  # 对应的MD文件路径
                        })
                    regions = formatted_regions
            except Exception:
                pass
            # Word：直接使用已导出的word_regions（已经是{Label, path}格式）
            try:
                word_regions = converted.get('word_regions')
                if word_regions:
                    regions = word_regions
            except Exception:
                pass
        
        # 注意：新的 word_regions 结构已经是 {Label, path} 格式，不需要额外合并
        
        # 6. 提取资源
        assets = []
        if extract_assets:
            assets = self._extract_assets(marked_content, regions, work_dir)
            # 合并全局Markdown图片到assets
            try:
                for img in global_md_images:
                    assets.append({
                        'type': 'image',
                        'name': img.get('name'),
                        'path': img.get('path'),
                        'source_region': None,
                        'format': img.get('format'),
                        'size': img.get('size')
                    })
            except Exception:
                pass
        
        # 7. 组装结果
        processing_time = 0.0  # 实际总耗时在分块等后续步骤完成后统一计算
        
        # 保存MD文件（仅为了保留原始内容，不用于分块）
        if converted['content_type'] == ContentType.MARKDOWN and marked_content:
            md_file_path = work_dir / f"{file_path.stem}.md"
            save_text(md_file_path, marked_content)
            logger.info(f"Markdown文件已保存: {md_file_path}")
            if str(md_file_path) not in markdown_files:
                markdown_files.append(str(md_file_path))
        
        # 确保markdown_files不为空（至少有一个主MD文件）
        if not markdown_files and file_type == FileType.EXCEL_XLSX:
            # 对于Excel，如果没有找到markdown文件，尝试从sheets/markdown目录查找
            md_dir = work_dir / "sheets" / "markdown"
            if md_dir.exists():
                for md_file in md_dir.glob("*.md"):
                    markdown_files.append(str(md_file))
        
        # ✅ 构建processing_info，包含file_id
        processing_info_dict = {
            **converted.get('metadata', {}),
            'processing_time': processing_time,
            'extract_regions': extract_regions,
            'extract_assets': extract_assets,
            'original_file': str(file_path.absolute()),  # 原始文件的绝对路径
            'markdown_files': markdown_files,  # 所有MD文件路径列表
        }
        
        # ✅ 如果有file_id，添加到processing_info
        if file_id:
            processing_info_dict['file_id'] = file_id
        
        result = PreprocessedDocument(
                source_file=str(file_path),
                file_type=file_type,
                content_type=converted['content_type'],
                text_content=converted.get('text', ''),
                markdown_content=marked_content if converted['content_type'] == ContentType.MARKDOWN else '',
                regions=regions,
                assets=assets,
                work_dir=work_dir,
                processing_info=processing_info_dict
            )
        
        # 8. 可选：对内容进行分块并落盘（默认启用；可用 CHUNKING_ENABLED=0 关闭）
        try:
            # ⚠️ 修复：支持对PLAIN_TEXT也进行分块
            content_for_chunking = None
            if result.content_type == ContentType.MARKDOWN and result.markdown_content:
                content_for_chunking = result.markdown_content
            elif result.content_type == ContentType.PLAIN_TEXT and result.text_content:
                # 对于PLAIN_TEXT，将其转换为简单的Markdown格式以便分块
                content_for_chunking = result.text_content
                logger.info("使用text_content进行分块（PLAIN_TEXT模式）")
            
            if content_for_chunking:
                enabled = str(os.getenv('CHUNKING_ENABLED', '1')).strip().lower() in ('1', 'true', 'yes', 'y', 'on')
                if enabled:
                    # 若已有结构化分块文件（sections），默认认为已分块，直接复用
                    existing_chunks = work_dir / f"{file_path.stem}_chunks_structured.json"
                    if existing_chunks.exists():
                        try:
                            # ✅ 保存相对路径（相对于项目根目录），跨平台兼容
                            try:
                                from pathlib import Path as P
                                # 尝试计算相对于当前工作目录的相对路径
                                rel_path = existing_chunks.relative_to(P.cwd())
                                result.processing_info['structured_chunks_file'] = str(rel_path).replace("\\", "/")
                            except ValueError:
                                # 如果无法计算相对于CWD的路径，直接使用文件名
                                result.processing_info['structured_chunks_file'] = f"{file_path.stem}_chunks_structured.json"
                            # 尝试读取 total_sections
                            data = json.loads(existing_chunks.read_text(encoding='utf-8'))
                            result.processing_info['chunks_total'] = int(data.get('total_sections') or 0)
                        except Exception:
                            pass
                        logger.info(f"检测到已有结构化分块文件，跳过重新分块: {existing_chunks}")
                    else:
                        # 检查分块模式
                        # 优先采用降级路径给出的 per-document 建议（无标题文档应走 character，避免 heading 退化为单块）
                        chunking_mode = (converted.get('metadata', {}) or {}).get('suggested_chunking_mode') or os.getenv("CHUNKING_MODE", "heading")  # character | heading (默认使用heading)
                        
                        # 初始化LLM服务（用于生成摘要；heading 与 character 两个分支共享）
                        from service.windows.preprocessing.preprocessing_function.heading_based_chunker import HeadingBasedChunker
                        llm_service = None
                        try:
                            from service.models import get_llm_service
                            llm_service = get_llm_service("validation")
                            logger.info("LLM服务初始化成功，将使用智能摘要")
                        except Exception as e:
                            logger.warning(f"无法初始化LLM服务，将使用简单摘要: {e}")

                        chunker = HeadingBasedChunker(llm_service=llm_service)

                        if chunking_mode == "heading":
                            # 使用基于一级标题的分块器
                            chunks_data = chunker.chunk_by_h1_headings(
                                content_for_chunking, 
                                file_path.name
                            )
                        else:
                            # character 模式：按段落感知打包分块，产出与 heading 一致的 sections 结构
                            # （短块<=CHAR_SUMMARY_DIRECT_CHARS 直接用内容作 summary；长块调 LLM 智能摘要）
                            chunks_data = chunker.chunk_by_paragraphs(
                                content_for_chunking,
                                file_path.name
                            )

                        # 保存结构化分块结果（heading 与 character 统一输出文件名与结构）
                        chunks_json_path = work_dir / f"{file_path.stem}_chunks_structured.json"
                        chunker.save_chunks_to_json(chunks_data, str(chunks_json_path))

                        result.processing_info['chunking_mode'] = chunking_mode
                        result.processing_info['chunks_total'] = chunks_data.get('total_sections') or len(chunks_data.get('sections', []))
                        # ✅ 保存相对路径（相对于项目根目录），跨平台兼容
                        try:
                            from pathlib import Path as P
                            rel_path = chunks_json_path.relative_to(P.cwd())
                            result.processing_info['structured_chunks_file'] = str(rel_path).replace("\\", "/")
                        except ValueError:
                            result.processing_info['structured_chunks_file'] = f"{file_path.stem}_chunks_structured.json"
                        result.processing_info['chunk_params'] = {'chunking_mode': chunking_mode}
        except Exception as e:
            logger.warning(f"Markdown 分块失败或未执行: {e}")

        # 统一计算总耗时（包含 OCR 兜底、分块、LLM 摘要等全部步骤）
        processing_time = time.time() - start_time
        result.processing_info['processing_time'] = processing_time

        # 9. 持久化标准化结果（preprocessed.json）
        try:
            self._save_preprocessed_package(result)
        except Exception as e:
            logger.warning(f"保存预处理标准包失败: {e}")

        logger.info(f"预处理完成: {processing_time:.2f}s, {len(regions)} 区域, {len(assets)} 资源")
        return result

    def _ensure_standard_regions(self, regions: list) -> list:
        """确保regions数组元素包含标准字段"""
        if not regions:
            return []
        
        standard_regions = []
        for region in regions:
            if isinstance(region, dict):
                standard_region = {
                    'Label': region.get('Label', ''),
                    'path': region.get('path', ''),
                }
                # markdown_path是可选的
                if 'markdown_path' in region:
                    standard_region['markdown_path'] = region['markdown_path']
                
                standard_regions.append(standard_region)
        
        return standard_regions
    
    def _ensure_standard_processing_info(self, processing_info: dict) -> dict:
        """确保processing_info包含所有标准字段"""
        standard_info = {
            'processing_time': processing_info.get('processing_time', 0.0),
            'extract_regions': processing_info.get('extract_regions', False),
            'extract_assets': processing_info.get('extract_assets', False),
            'original_file': processing_info.get('original_file', ''),
            'markdown_files': processing_info.get('markdown_files', []),
        }
        
        # 可选字段（如果存在则添加）
        optional_fields = [
            'chunks_total',
            'structured_chunks_file',
            'chunking_mode',
            'chunk_params',
            'conversion_method',
            'steps',
            'sheets_total',
        ]
        
        for field in optional_fields:
            if field in processing_info:
                standard_info[field] = processing_info[field]
        
        return standard_info

    def _save_preprocessed_package(self, pre: PreprocessedDocument) -> None:
        """将标准化数据包保存为 JSON，供主流程只读使用。"""
        try:
            import json
            out_dir = pre.work_dir or (self.cache_dir / Path(pre.source_file).stem)
            out_dir = Path(out_dir)
            out_dir.mkdir(parents=True, exist_ok=True)
            
            # 根据文件类型决定是否标准化
            # Word/PDF需要分块相关字段
            if pre.file_type in [FileType.WORD_DOCX, FileType.WORD_DOC, FileType.PDF]:
                # Word/PDF保持原有结构，确保分块相关字段
                pkg = {
                    'source_file': pre.source_file,
                    'file_type': pre.file_type.value,
                    'content_type': pre.content_type.value,
                    'text_content': pre.text_content,  # Word/PDF可能需要文本内容
                    'markdown_content': pre.markdown_content,  # Word/PDF需要markdown内容
                    'regions': pre.regions or [],
                    'assets': pre.assets or [],
                    'processing_info': pre.processing_info or {},
                    'work_dir': str(out_dir),
                }
            # Excel/RTF使用相同结构
            elif pre.file_type in [FileType.EXCEL_XLSX, FileType.RTF]:
                # Excel/RTF标准化regions
                standardized_regions = self._ensure_standard_regions(pre.regions or [])
                standardized_info = self._ensure_standard_processing_info(pre.processing_info or {})
                
                pkg = {
                    'source_file': pre.source_file,
                    'file_type': pre.file_type.value,
                    'content_type': pre.content_type.value,
                    'regions': standardized_regions,
                    'assets': pre.assets or [],
                    'processing_info': standardized_info,
                    'work_dir': str(out_dir),
                }
            else:
                # 其他文件类型保持原样
                pkg = {
                    'source_file': pre.source_file,
                    'file_type': pre.file_type.value,
                    'content_type': pre.content_type.value,
                    'regions': pre.regions or [],
                    'assets': pre.assets or [],
                    'processing_info': pre.processing_info or {},
                    'work_dir': str(out_dir),
                }
            
            # 仅当structured_data不为空时添加（向后兼容）
            if pre.structured_data:
                pkg['structured_data'] = pre.structured_data
            
            # 仅在分块文件存在时添加chunks_file字段
            if 'structured_chunks_file' in pre.processing_info:
                chunks_path = Path(pre.processing_info['structured_chunks_file'])
                if chunks_path.exists():
                    pkg['chunks_file'] = chunks_path.name
            preprocessed_json_path = out_dir / 'preprocessed.json'
            preprocessed_json_path.write_text(
                json.dumps(pkg, ensure_ascii=False, indent=2), encoding='utf-8')
            logger.info(f"预处理标准包已保存: {preprocessed_json_path}")
            
            # ✅ 更新文件映射 (file_mappings.json)
            try:
                from service.windows.preprocessing.preprocessing_function.file_mappings_manager import add_file_mapping
                
                # 从source_file提取项目名称和文件名
                source_path = Path(pre.source_file)
                file_name = source_path.name
                
                # 尝试从路径推断项目名称
                # 路径格式: AAA/project_data/{project_name}/{category}/{filename}
                project_name = "default"
                try:
                    parts = source_path.parts
                    if 'project_data' in parts:
                        idx = parts.index('project_data')
                        if idx + 1 < len(parts):
                            project_name = parts[idx + 1]
                except Exception:
                    pass
                
                # 从processing_info获取file_id（如果有）
                file_id = pre.processing_info.get('file_id', '') or str(hash(file_name))[:16]
                
                # 添加到映射
                add_file_mapping(
                    project_name=project_name,
                    file_name=file_name,
                    file_id=file_id,
                    preprocessed_json_path=str(preprocessed_json_path),
                    preprocessed_dir=str(out_dir),
                    status="success"
                )
                logger.info(f"✓ 文件映射已更新: {project_name}/{file_name}")
            except Exception as mapping_err:
                logger.warning(f"更新文件映射失败: {mapping_err}")
            
            # 更新预处理索引（保留旧的索引逻辑以兼容）
            try:
                from service.windows.preprocessing.preprocessing_function.preprocessing_index import add_preprocessing_result
                
                file_name = Path(pre.source_file).name
                
                # 查找chunks_file路径
                chunks_file = None
                if 'structured_chunks_file' in pre.processing_info:
                    chunks_file = str(out_dir / pre.processing_info['structured_chunks_file'])
                
                # 查找regions目录
                regions_dir = None
                if pre.regions:
                    for region in pre.regions:
                        if isinstance(region, dict):
                            label = region.get('Label', '')
                            if label.startswith('Table_') or label.startswith('Image_'):
                                # Word/PDF类型
                                regions_dir = str(out_dir / 'regions_word')
                                break
                            elif label.startswith('Sheet_'):
                                # Excel/RTF类型
                                regions_dir = str(out_dir / 'sheets')
                                break
                
                # 添加到索引
                add_preprocessing_result(
                    file_name=file_name,
                    original_path=pre.source_file,
                    preprocessing_dir=str(out_dir),
                    preprocessed_json=str(out_dir / 'preprocessed.json'),
                    file_type=pre.file_type.value,
                    chunks_file=chunks_file,
                    regions_dir=regions_dir,
                    status="success"
                )
                logger.info(f"📋 预处理索引已更新: {file_name}")
            except Exception as idx_err:
                logger.warning(f"更新预处理索引失败: {idx_err}")
                
        except Exception as e:
            logger.warning(f"写入预处理标准包失败: {e}")
    
    def _detect_file_type(self, file_path: Path, force_ocr: bool) -> FileType:
        """检测文件类型"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return FileType.WORD_DOCX
        elif suffix == '.doc':
            return FileType.WORD_DOC
        elif suffix == '.rtf':
            # RTF文件走独立的RTF Pipeline（RTF→Excel→拆分Sheet→Markdown）
            return FileType.RTF
        elif suffix == '.xlsx':
            return FileType.EXCEL_XLSX
        elif suffix == '.xls':
            return FileType.EXCEL_XLS
        elif suffix == '.pdf':
            return FileType.PDF
        else:
            return FileType.UNKNOWN

    def _normalize_format(self, file_path: Path, file_type: FileType, work_dir: Path) -> Path:
        """
        使用 LibreOffice 将 .doc/.xls 转为 .docx/.xlsx，
        使后续流程统一使用 python-docx / openpyxl 处理，彻底消除 COM 依赖。
        """
        import subprocess

        if file_type == FileType.WORD_DOC:
            target_fmt = 'docx'
        elif file_type == FileType.EXCEL_XLS:
            target_fmt = 'xlsx'
        else:
            return file_path

        converted = work_dir / f"{file_path.stem}.{target_fmt}"

        # 复用已有转换结果
        if converted.exists() and converted.stat().st_size > 0:
            logger.info(f"复用已有格式转换: {converted}")
            return converted

        try:
            with _LIBREOFFICE_LOCK:
                subprocess.run([
                    'soffice', '--headless', '--convert-to', target_fmt,
                    '--outdir', str(work_dir), str(file_path)
                ], check=True, timeout=120, capture_output=True)

            if converted.exists() and converted.stat().st_size > 0:
                logger.info(f"格式转换成功: {file_path.name} → {converted.name}")
                return converted
            else:
                raise RuntimeError("LibreOffice 未生成目标文件")
        except subprocess.TimeoutExpired:
            raise RuntimeError(f"LibreOffice 转换超时: {file_path.name}")
        except Exception:
            raise RuntimeError(
                f"格式转换失败: {file_path.name}，请确保 LibreOffice 已安装且在 PATH 中"
            )

    def _normalize_content_type(self, content_type: Optional[str]) -> Optional[str]:
        """
        规范化内容类型：将用户输入的类型映射到内部使用的类型
        
        映射关系：
        - 'txt' -> 'text' (文本类型)
        - 'jpg' -> 'images' (纯图类型)
        - 'list' -> 'tables' (纯表类型)
        - 也支持直接使用内部类型 ('text', 'images', 'tables')
        """
        if not content_type:
            return None
        content_type_lower = content_type.lower().strip()
        mapping = {
            'txt': 'text',
            'jpg': 'images',
            'jpeg': 'images',
            'png': 'images',
            'image': 'images',
            'img': 'images',
            'list': 'tables',
            'table': 'tables',
            'tbl': 'tables',
        }
        # 如果已经在映射表中，返回映射后的值
        if content_type_lower in mapping:
            return mapping[content_type_lower]
        # 如果已经是内部类型，直接返回
        if content_type_lower in ('text', 'images', 'tables'):
            return content_type_lower
        # 未知类型，返回None（将使用自动分类）
        logger.warning(f"未知的内容类型: {content_type}，将使用自动分类")
        return None

    def _convert_to_intermediate(self, file_path: Path, file_type: FileType, work_dir: Path, content_type_override: Optional[str] = None) -> Dict[str, Any]:
        """转换为中间格式"""
        if file_type == FileType.RTF:
            # RTF文档 → 独立的RTF Pipeline（RTF→Excel→拆分Sheet→Markdown）
            try:
                from service.windows.preprocessing.preprocessing_function.rtf.rtf_pipeline import rtf_run
                return rtf_run(file_path, work_dir)
            except Exception as e:
                logger.error(f"RTF Pipeline失败: {e}", exc_info=True)
                raise RuntimeError(f"RTF转换失败: {e}")
        
        elif file_type == FileType.WORD_DOCX:
            # Word文档（.docx）
            # 基于内容类型分类（文本/纯图片/纯表）选择最佳路径
            normalized_type = self._normalize_content_type(content_type_override)
            # 委派到 Word Pipeline
            try:
                from service.windows.preprocessing.preprocessing_function.word.word_pipeline import word_run as word_run_docx
                return word_run_docx(file_path, work_dir, mode=normalized_type)
            except Exception as e:
                logger.warning(f"Word(.docx) 管道失败，回退到内置流程: {e}")
                return self._convert_word_by_kind(file_path, work_dir, content_type_override=normalized_type)
        
        elif file_type == FileType.WORD_DOC:
            # .doc → LibreOffice 转 .docx → 走 docx pipeline
            docx_path = self._normalize_format(file_path, file_type, work_dir)
            normalized_type = self._normalize_content_type(content_type_override)
            try:
                from service.windows.preprocessing.preprocessing_function.word.word_pipeline import word_run as word_run_docx
                return word_run_docx(docx_path, work_dir, mode=normalized_type)
            except Exception as e:
                logger.warning(f"Word(.doc) 管道失败，回退到内置流程: {e}")
                return self._convert_word_by_kind(docx_path, work_dir, content_type_override=normalized_type)
        
        elif file_type == FileType.PDF:
            # 正常PDF → 直接转 Markdown（OCR 优先，PyMuPDF 兜底）
            try:
                from service.windows.preprocessing.preprocessing_function.pdf.pdf_pipeline import pdf_run
                return pdf_run(file_path, work_dir, scanned=False)
            except Exception as e:
                logger.error(f"PDF Pipeline失败: {e}", exc_info=True)
                raise RuntimeError(f"PDF转换失败: {e}")

        elif file_type == FileType.EXCEL_XLS:
            # .xls → LibreOffice 转 .xlsx → 走 xlsx pipeline
            xlsx_path = self._normalize_format(file_path, file_type, work_dir)
            try:
                from service.windows.preprocessing.preprocessing_function.excel.excel_pipeline import excel_run
                return excel_run(xlsx_path, work_dir)
            except Exception as e:
                logger.error(f"Excel(.xls) 管道失败: {e}", exc_info=True)
                raise RuntimeError(f"Excel转换失败: {e}")

        elif file_type == FileType.EXCEL_XLSX:
            # 委派到 Excel Pipeline：拆分Sheet并生成Markdown
            try:
                from service.windows.preprocessing.preprocessing_function.excel.excel_pipeline import excel_run
                return excel_run(file_path, work_dir)
            except Exception as e:
                logger.error(f"Excel 管道失败: {e}", exc_info=True)
                raise RuntimeError(f"Excel转换失败: {e}")
        
        else:
            # 未知格式 → 尝试读取纯文本
            try:
                content = file_path.read_text(encoding='utf-8')
                return {
                    'content': content,
                    'content_type': ContentType.PLAIN_TEXT,
                    'text': content,
                    'metadata': {'conversion_method': 'fallback_text'}
                }
            except Exception:
                return {
                    'content': '',
                    'content_type': ContentType.PLAIN_TEXT,
                    'text': '',
                    'metadata': {'conversion_method': 'failed'}
                }
    
    def _word_to_markdown(self, word_path: Path, work_dir: Path) -> Dict[str, Any]:
        """
        Word转Markdown完整流程（标准流程）
        
        流程：
        1. Word → 标记表格/图片（插入{{Table_1_Start}}标签）
        2. 并行执行：
           a) 标记后的Word → 直接导出表格/图片为独立DOCX文件
           b) 标记后的Word → PDF → Markdown（OCR）
        3. 返回Markdown + 导出的区域文件
        """
        try:
            logger.info(f"开始Word完整转换流程: {word_path}")
            # 复用缓存：若存在 md 与 manifest，直接读取
            # 使用原Word文件名命名MD文件
            md_file = work_dir / f"{word_path.stem}.md"
            stage = self._stage_paths(work_dir, "word_md")
            if md_file.exists() and stage["manifest"].exists():
                try:
                    manifest = json.loads(stage["manifest"].read_text(encoding="utf-8"))
                    markdown_content = md_file.read_text(encoding="utf-8")
                    marked_word = manifest.get("marked_word")
                    pdf_file = manifest.get("pdf_file")
                    cached_regions = manifest.get("word_regions") or []
                    logger.info("复用缓存的 Word→Markdown 结果")
                    return {
                        'content': markdown_content,
                        'content_type': ContentType.MARKDOWN,
                        'text': markdown_content,
                        'word_regions': cached_regions,
                        'metadata': {
                            'conversion_method': 'word_mark_pdf_ocr',
                            'marked_word': str(marked_word) if marked_word else None,
                            'pdf_file': str(pdf_file) if pdf_file else None,
                            'steps': ['mark_word', 'export_regions', 'word_to_pdf', 'pdf_to_md'],
                            'cache_reused': True,
                            'word_regions_count': len(cached_regions)
                        }
                    }
                except Exception:
                    pass
            
            # Step 1: 标记Word文档中的表格和图片
            from service.windows.preprocessing.preprocessing_function.word.word_pipeline import word_marker
            marked_word = word_marker.mark_tables_and_images(word_path, work_dir)
            if not marked_word or not marked_word.exists():
                logger.warning("Word标记失败，使用原文件")
                marked_word = word_path
            
            # Step 2a: 直接从标记的Word导出表格和图片区域（不需要等待Markdown转换）
            word_regions: List[Dict[str, Any]] = []
            try:
                exported_regions = self._export_word_regions(marked_word, work_dir)
                for region_name, region_path in exported_regions.items():
                    word_regions.append({
                        'Label': region_name,
                        'path': region_path
                    })
                logger.info(f"直接导出 {len(word_regions)} 个Word区域（不依赖Markdown转换）")
            except Exception as e:
                logger.warning(f"导出Word区域失败: {e}")
            
            # Step 2b: 标记后的Word → PDF
            pdf_file = self._word_to_pdf(marked_word, work_dir)
            if not pdf_file or not pdf_file.exists():
                logger.error("Word→PDF转换失败")
                if self._ocr_strict():
                    raise RuntimeError("OCR_STRICT=1: Word→PDF failed")
                return self._fallback_word_text_extraction(word_path, marked_word_path=marked_word)
            
            # Step 3: PDF → Markdown（OCR）
            from service.windows.preprocessing.preprocessing_function.pdf.pdf_pipeline import pdf_to_markdown_ocr
            markdown_content = pdf_to_markdown_ocr(pdf_file, work_dir)
            if not markdown_content:
                logger.error("PDF→Markdown失败")
                if self._ocr_strict():
                    raise RuntimeError("OCR_STRICT=1: PDF→Markdown failed")
                return self._fallback_word_text_extraction(word_path, marked_word_path=marked_word, pdf_path=pdf_file)
            logger.info(f"Word转换成功，Markdown长度: {len(markdown_content)} 字符")
            
            # 写入缓存文件与manifest（包括word_regions）
            try:
                save_text(md_file, markdown_content)
                manifest = {
                    "marked_word": str(marked_word),
                    "pdf_file": str(pdf_file),
                    "markdown_file": str(md_file),
                    "word_regions": word_regions,
                    "created_at": __import__("datetime").datetime.now().isoformat(timespec='seconds')
                }
                save_json(stage["manifest"], manifest)
                stage["done"].write_text("ok", encoding="utf-8")
            except Exception:
                pass
            
            return {
                'content': markdown_content,
                'content_type': ContentType.MARKDOWN,
                'text': markdown_content,
                'word_regions': word_regions,
                'metadata': {
                    'conversion_method': 'word_mark_pdf_ocr',
                    'marked_word': str(marked_word),
                    'pdf_file': str(pdf_file),
                    'steps': ['mark_word', 'export_regions', 'word_to_pdf', 'pdf_to_md'],
                    'word_regions_count': len(word_regions)
                }
            }
            
        except Exception as e:
            if self._ocr_strict():
                logger.error(f"Word完整转换失败(strict): {e}")
                raise
            logger.error(f"Word完整转换失败: {e}，回退到简单提取", exc_info=True)
            return self._fallback_word_text_extraction(word_path)

    def _convert_word_by_kind(self, word_path: Path, work_dir: Path, content_type_override: Optional[str] = None) -> Dict[str, Any]:
        """
        根据 Word 内容类型（文本/纯图片/纯表）选择转换策略：
        - 文本：在 Word 内标记 → Word→PDF→OCR→Markdown（分块）
        - 纯图片：不处理（跳过，留待视觉模型后续处理）
        - 纯表：仅做拆分导出，每个表格生成独立 .docx
        
        Args:
            word_path: Word文件路径（包括.docx, .doc, .rtf）
            work_dir: 工作目录
            content_type_override: 手动指定的内容类型（'text', 'images', 'tables'），如果提供则覆盖自动分类
        """
        # 如果提供了手动指定的类型，使用它；否则自动分类
        if content_type_override:
            kind = content_type_override
            logger.info(f"使用手动指定的Word内容类型: {kind}")
        else:
            try:
                kind = self._classify_word_content(word_path)
                logger.info(f"自动分类Word内容类型: {kind}")
            except Exception as e:
                kind = "text"  # 无法分类时走稳健路径（PDF+OCR）
                logger.warning(f"Word内容分类失败，{e}，使用默认类型: {kind}")

        if kind == "tables":
            # 仅做表格拆分：每个表格导出为独立的 .docx
            return self._word_tables_split_to_docx(word_path, work_dir)
        if kind == "images":
            # 仅导出图片：把所有图片提取到 images/ 文件夹
            return self._word_images_extract_to_files(word_path, work_dir)
        else:
            # 其余（文本为主）：按原本Word路径（标记→PDF→OCR→Markdown）
            out = self._word_to_markdown(word_path, work_dir)
            try:
                if content_type_override:
                    out.setdefault('metadata', {})['content_type_override'] = content_type_override
            except Exception:
                pass
            return out

    def _classify_word_content(self, word_path: Path) -> str:
        """
        粗粒度分类 Word 内容：'text' | 'images' | 'tables'
        规则（启发式）：
        - 若表格单元格数量大且段落文字极少 → tables
        - 若存在图片且段落文字极少且表格很少 → images
        - 否则 → text
        """
        try:
            from docx import Document  # type: ignore
            doc = Document(word_path)
        except Exception:
            return "text"  # 无法打开，走稳健路径

        try:
            text_chars = sum(len((p.text or '').strip()) for p in doc.paragraphs)
        except Exception:
            text_chars = 0

        try:
            tables_count = len(doc.tables)
            table_cells = 0
            for t in doc.tables:
                try:
                    for row in t.rows:
                        table_cells += len(row.cells)
                except Exception:
                    pass
        except Exception:
            tables_count = 0
            table_cells = 0

        # 统计图片数量（通过底层XML blip/drawing/pict）
        images_count = 0
        try:
            body = doc._element.body
            for element in list(body):
                tag = getattr(element, 'tag', '') or ''
                if tag.endswith('p'):
                    has_image = False
                    for child in element.iter():
                        ctag = getattr(child, 'tag', '') or ''
                        if ctag.endswith('blip') or ctag.endswith('drawing') or ctag.endswith('pict'):
                            has_image = True
                            break
                    if has_image:
                        images_count += 1
        except Exception:
            images_count = 0

        # 启发式阈值
        low_text = text_chars < 200
        many_table_cells = table_cells >= 30 or tables_count >= 2
        has_images = images_count >= 1

        if many_table_cells and low_text:
            return "tables"
        if has_images and low_text and not many_table_cells:
            return "images"
        return "text"

    def _table_to_markdown(self, table, index: int) -> str:
        """将单个 docx 表格序列化为 Markdown 表格（含 {{Table_N_Start/End}} 标签）。
        与 heading_based_chunker 的 _TABLE_START_RE/_TABLE_END_RE 约定一致。"""
        rows = []
        try:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    # 合并换行，清理竖线
                    cell_text = (cell.text or '').replace('\n', ' ').replace('|', '\\|').strip()
                    cells.append(cell_text)
                rows.append(cells)
        except Exception:
            return ""

        if not rows:
            return ""

        # 构造 Markdown 表格（首行作表头，其余为正文行）
        header = rows[0]
        body_rows = rows[1:] if len(rows) > 1 else []
        parts = [f"{{{{Table_{index}_Start}}}}"]
        parts.append("| " + " | ".join(header) + " |")
        parts.append("| " + " | ".join(["---"] * len(header)) + " |")
        for r in body_rows:
            parts.append("| " + " | ".join(r) + " |")
        parts.append(f"{{{{Table_{index}_End}}}}")
        return "\n".join(parts)

    def _heading_level(self, paragraph) -> "Optional[int]":
        """从段落样式判断 markdown 标题层级（1-6），非标题返回 None。
        兼容英文 'Heading N'、中文本地化 '标题 N'、'Title'/'标题'、style_id='HeadingN'。"""
        import re
        try:
            name = paragraph.style.name
        except Exception:
            name = None
        try:
            style_id = paragraph.style.style_id
        except Exception:
            style_id = None

        if name:
            m = re.match(r'^(?:heading|标题)\s*(\d+)$', name.strip(), re.IGNORECASE)
            if m:
                return min(int(m.group(1)), 6)
            if name.strip().lower() in ('title', '标题'):
                return 1
        if style_id:
            m = re.match(r'^Heading(\d+)$', str(style_id))
            if m:
                return min(int(m.group(1)), 6)
        return None

    def _word_to_pdf(self, word_path: Path, work_dir: Path) -> Optional[Path]:
        """
        将Word文档转换为PDF格式。
        使用 LibreOffice 通过命令行工具转换
        
        Args:
            word_path: 源Word文档的路径
            work_dir: 工作目录路径，用于存放生成的PDF文件
        
        Returns:
            成功时返回生成PDF文件的路径，失败时返回None
            如果PDF已存在（上一次转换结果），则直接返回该路径
        """
        try:
            pdf_file = work_dir / f"{word_path.stem}.pdf"
            # 若已有PDF，直接复用
            try:
                if pdf_file.exists():
                    return pdf_file
            except Exception:
                pass
            # 使用LibreOffice对文件格式进行转换
            try:
                import subprocess
                with _LIBREOFFICE_LOCK:
                    subprocess.run([
                        'soffice', '--headless', '--convert-to', 'pdf',
                        '--outdir', str(work_dir), str(word_path)
                    ], check=True, timeout=60, capture_output=True)
                
                # 验证转换结果是否生成成功
                if pdf_file.exists():
                    logger.info(f"Word→PDF成功（LibreOffice）: {pdf_file}")
                    return pdf_file
            except Exception as e:
                logger.warning(f"LibreOffice转PDF失败: {e}")
            
            logger.error("所有Word→PDF方法均失败")
            return None
            
        except Exception as e:
            logger.error(f"Word→PDF失败: {e}", exc_info=True)
            return None
    
    def _fallback_word_text_extraction(self, word_path: Path, marked_word_path: Path = None, pdf_path: Path = None) -> Dict[str, Any]:
        """
        回退方案：用 python-docx 直接提取 Word 内容并构造结构化 Markdown（不经过 PDF/OCR）。
        - 段落按样式补 # 标题，保留文档真实顺序
        - 表格按 {{Table_N}} markdown 表格输出，与 HeadingBasedChunker 的表格约定一致
        - 无标题样式的文档标记 suggested_chunking_mode='character'，避免 heading 分块退化为单块

        Args:
            word_path: 原始Word文档路径
            marked_word_path: 标记后的Word文档路径（如果存在），保留此路径以便后续导出区域
            pdf_path: 转换后的PDF路径（如果存在）
        """
        try:
            from docx import Document  # type: ignore
            from docx.text.paragraph import Paragraph  # type: ignore
            from docx.table import Table  # type: ignore
            doc = Document(word_path)

            md_parts: List[str] = []
            has_headings = False
            table_index = 0

            # 按文档真实顺序遍历段落与表格（doc.paragraphs + doc.tables 会丢顺序）
            for child in doc.element.body:
                tag = child.tag
                if tag.endswith('}p'):
                    p = Paragraph(child, doc)
                    text = (p.text or '').strip()
                    level = self._heading_level(p)
                    if level is not None:
                        has_headings = True
                        if text:
                            md_parts.append(f"{'#' * level} {text}")
                    else:
                        if text:
                            md_parts.append(text)
                elif tag.endswith('}tbl'):
                    table_index += 1
                    table_md = self._table_to_markdown(Table(child, doc), table_index)
                    if table_md:
                        md_parts.append(table_md)

            md_content = "\n\n".join(md_parts).strip()
            logger.info(f"回退提取Word结构化Markdown: {len(md_content)} 字符, has_headings={has_headings}, tables={table_index}")

            metadata = {
                'conversion_method': 'fallback_docx_structured_md',
                'has_headings': has_headings,
                'suggested_chunking_mode': 'heading' if has_headings else 'character',
                'tables_count': table_index,
            }
            word_regions: List[Dict[str, Any]] = []
            
            # ⚠️ 关键修复：即使走回退路径，也要导出regions
            if marked_word_path and Path(marked_word_path).exists():
                metadata['marked_word'] = str(marked_word_path)
                logger.info(f"保留标记Word路径并导出区域: {marked_word_path}")
                
                try:
                    # 确定work_dir（从marked_word_path的父目录）
                    work_dir = Path(marked_word_path).parent
                    exported_regions = self._export_word_regions(Path(marked_word_path), work_dir)
                    for region_name, region_path in exported_regions.items():
                        word_regions.append({
                            'Label': region_name,
                            'path': region_path
                        })
                    logger.info(f"✅ 回退路径导出了 {len(word_regions)} 个regions")
                except Exception as e:
                    logger.warning(f"回退路径导出regions失败: {e}")
            
            if pdf_path:
                metadata['pdf_file'] = str(pdf_path)
            
            metadata['word_regions_count'] = len(word_regions)
            
            return {
                'content': md_content,
                'content_type': ContentType.MARKDOWN,
                'text': md_content,
                'word_regions': word_regions,  # ⚠️ 添加word_regions
                'metadata': metadata
            }
        except Exception as e:
            logger.error(f"回退提取也失败: {e}", exc_info=True)
            return {
                'content': '',
                'content_type': ContentType.PLAIN_TEXT,
                'text': '',
                'word_regions': [],
                'metadata': {'conversion_method': 'failed', 'error': str(e)}
            }
    
    def _extract_assets(self, content: str, regions: List[Dict], work_dir: Path) -> List[Dict[str, Any]]:
        """提取资源（图片等）"""
        assets = []
        # 从regions中收集已提取的图片
        for region in regions:
            for img in region.get('images', []):
                assets.append({
                    'type': 'image',
                    'name': img['name'],
                    'path': img['path'],
                    'source_region': region['name'],
                    'format': img.get('format'),
                    'size': img.get('size')
                })
        return assets

    def _export_word_regions(self, marked_word: Path, work_dir: Path) -> Dict[str, str]:
        """
        将已标记的Word文档中的各标签区间导出为独立的Word文件。
        返回 {region_name: exported_docx_path}
        
        使用简单可靠的 python-docx 方法，直接复制标签之间的内容
        """
        try:
            from service.windows.preprocessing.preprocessing_function.word.word_pipeline import word_region_extractor
            export_dir = work_dir / 'regions_word'
            export_dir.mkdir(parents=True, exist_ok=True)
            
            # 使用新的简单提取器
            exported = word_region_extractor.extract_regions(
                marked_file=str(marked_word),
                export_dir=str(export_dir)
            ) or []
            
            logger.info(f"成功导出 {len(exported)} 个Word区域")
            return {item.get('name'): item.get('path') for item in exported if item.get('name') and item.get('path')}
        except Exception as e:
            logger.warning(f"导出Word标签区间发生异常: {e}")
            import traceback
            logger.warning(traceback.format_exc())
            return {}

    def _word_tables_split_to_docx(self, word_path: Path, work_dir: Path) -> Dict[str, Any]:
        """
        将 Word 文档中的每个表格拆分导出为独立的 .docx 文件（不做OCR/Markdown）。
        返回结构化结果，并提供 word_regions 列表（仅包含表格）。
        """
        try:
            # 复用缓存
            sp = self._stage_paths(work_dir, "word_tables_split")
            if sp["manifest"].exists():
                try:
                    manifest = json.loads(sp["manifest"].read_text(encoding="utf-8"))
                    regions = manifest.get("regions") or []
                    if regions:
                        logger.info("复用缓存的 Word 表格拆分结果")
                        return {
                            'content': '',
                            'content_type': ContentType.STRUCTURED,
                            'text': '',
                            'word_regions': regions,
                            'metadata': {
                                'conversion_method': 'word_tables_split_docx',
                                'tables_total': len(regions),
                                'cache_reused': True
                            }
                        }
                except Exception:
                    pass
            from service.windows.insertion.word_document_service import word_document_service
            export_dir = work_dir / 'regions_word'
            export_dir.mkdir(parents=True, exist_ok=True)
            exported = word_document_service.export_all_tables_and_images(str(word_path), str(export_dir)) or []
            regions: List[Dict[str, Any]] = []
            table_idx = 0
            image_idx = 0
            for item in exported:
                name = item.get('name') or ''
                path = item.get('path') or ''
                itype = (item.get('type') or '').lower()
                if not name or not path:
                    continue
                # 使用新的 {Label, path} 结构
                if itype == 'table' or name.startswith('Table_'):
                    table_idx += 1
                elif itype in ('image', 'chart', 'shape') or name.startswith('Image_'):
                    image_idx += 1
                
                regions.append({
                    "Label": name,
                    "path": path
                })
            result = {
                'content': '',
                'content_type': ContentType.STRUCTURED,
                'text': '',
                'word_regions': regions,
                'metadata': {
                    'conversion_method': 'word_tables_split_docx',
                    'tables_total': table_idx,
                    'images_total': image_idx,
                    'objects_total': len(regions)
                }
            }
            # 写入manifest
            try:
                save_json(sp["manifest"], {"regions": regions})
                sp["done"].write_text("ok", encoding="utf-8")
            except Exception:
                pass
            return result
        except Exception as e:
            logger.error(f"Word表格拆分导出失败，回退到标准Markdown转换: {e}", exc_info=True)
            return self._word_to_markdown(word_path, work_dir)

    def _word_images_extract_to_files(self, word_path: Path, work_dir: Path) -> Dict[str, Any]:
        """
        将 Word 文档中的所有图片提取到 images/ 目录下（直接导出为图片文件）。
        返回结构化结果，并提供 word_regions 列表（每张图片一个条目）。
        """
        try:
            # 复用缓存
            sp = self._stage_paths(work_dir, "word_images_extract")
            if sp["manifest"].exists():
                try:
                    manifest = json.loads(sp["manifest"].read_text(encoding="utf-8"))
                    extracted = manifest.get("regions") or []
                    images_dir = manifest.get("images_dir")
                    if extracted and images_dir:
                        logger.info("复用缓存的 Word 图片提取结果")
                        return {
                            'content': '',
                            'content_type': ContentType.STRUCTURED,
                            'text': '',
                            'word_regions': extracted,
                            'metadata': {
                                'conversion_method': 'word_images_extract',
                                'images_total': len(extracted),
                                'images_dir': images_dir,
                                'cache_reused': True
                            }
                        }
                except Exception:
                    pass
            images_dir = work_dir / 'images'
            images_dir.mkdir(parents=True, exist_ok=True)
            # 确保是docx：若是doc/rtf已在上游转换/分流，这里按docx zip方式提取
            from zipfile import ZipFile
            extracted: List[Dict[str, Any]] = []
            try:
                with ZipFile(str(word_path), 'r') as zf:
                    # DOCX 中图片位于 word/media/*
                    media_files = [n for n in zf.namelist() if n.startswith('word/media/')]
                    media_files.sort()
                    count = 0
                    for n in media_files:
                        # 提取到 images_dir，命名按顺序
                        suffix = ''
                        try:
                            from pathlib import PurePosixPath
                            suffix = PurePosixPath(n).suffix
                        except Exception:
                            pass
                        count += 1
                        out_path = images_dir / f"Image_{count:03d}{suffix or ''}"
                        try:
                            with zf.open(n) as src, open(out_path, 'wb') as dst:
                                dst.write(src.read())
                            extracted.append({
                                "name": f"Image_{count}",
                                "type": "image",
                                "index": count,
                                "file_path": str(out_path)
                            })
                        except Exception:
                            count -= 1
                            continue
            except Exception as e:
                logger.error(f"DOCX图片解包失败，回退到标准Markdown转换: {e}", exc_info=True)
                return self._word_to_markdown(word_path, work_dir)
            result = {
                'content': '',
                'content_type': ContentType.STRUCTURED,
                'text': '',
                'word_regions': extracted,
                'metadata': {
                    'conversion_method': 'word_images_extract',
                    'images_total': len(extracted),
                    'images_dir': str(images_dir)
                }
            }
            try:
                save_json(sp["manifest"], {"regions": extracted, "images_dir": str(images_dir)})
                sp["done"].write_text("ok", encoding="utf-8")
            except Exception:
                pass
            return result
        except Exception as e:
            logger.error(f"Word图片提取失败，回退到标准Markdown转换: {e}", exc_info=True)
            return self._word_to_markdown(word_path, work_dir)
