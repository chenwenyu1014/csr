# -*- coding: utf-8 -*-
from __future__ import annotations
import logging
import re
import shutil
import copy
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple, Set
from docx import Document
from lxml import etree
import zipfile
import tempfile

# Word XML 命名空间
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WNS_P = f'{{{WNS}}}'

# DrawingML 和 Relationships 命名空间（用于扫描引用的媒体文件）
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'

# 已压缩的图片扩展名，打包时用 STORED 避免重复压缩
_MEDIA_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.emf', '.wmf'}


from service.windows.preprocessing.file_processor import FileProcessor

logger = logging.getLogger(__name__)


def word_run(word_path: Path | str, work_dir: Path | str, mode: Optional[str] = None) -> Dict[str, Any]:
    """
    Word (.docx/.rtf-as-docx-class) pipeline.
    mode: Optional logical hint ('text'|'images'|'tables'), currently passed through to existing logic.
    """
    fp = FileProcessor()
    word_path = Path(word_path)
    work_dir = Path(work_dir)
    # Delegate to existing robust path
    normalized = mode
    return fp._convert_word_by_kind(word_path, work_dir, content_type_override=normalized)  # type: ignore


class WordRegionMarker:
    """
    在 Word 文档的表格与图片前后插入标记段落。

    通过解析 document.xml，在顶层表格、顶层图片段落，以及表格单元格内的
    嵌套表格和图片前后插入 {{Table_N_Start}} / {{Table_N_End}} 等标记，
    供 WordRegionExtractor 按区域裁剪为独立文档。
    """

    # 需要标记为图片的 OLE 对象 ProgID 列表（排除 Excel 等表格类对象）
    _IMAGE_OLE_PROGIDS = frozenset([
        'Visio.Drawing',      # Visio 绘图
        'MSVisio',            # Visio 的另一种 ProgID 格式
        'PowerPoint.Show',    # PowerPoint 幻灯片
        'PowerPoint.Slide',   # PowerPoint 幻灯片
        'Equation',           # 公式编辑器
        'WordArt',            # 艺术字
    ])

    def _create_marker_paragraph_element(self, marker_text: str) -> 'etree._Element':
        """
        创建包含标记文本的段落 XML 元素（不添加到文档）。

        Args:
            marker_text: 标记文本，如 "{{Table_1_Start}}"

        Returns:
            段落 XML 元素 (w:p)
        """
        from lxml import etree

        p = etree.Element(f'{{{WNS}}}p')

        pPr = etree.SubElement(p, f'{{{WNS}}}pPr')
        jc = etree.SubElement(pPr, f'{{{WNS}}}jc')
        jc.set(f'{{{WNS}}}val', 'center')

        r = etree.SubElement(p, f'{{{WNS}}}r')
        t = etree.SubElement(r, f'{{{WNS}}}t')
        t.text = marker_text

        return p

    def _find_ref_paragraph_before(self, elements: list, target_element):
        """
        从目标元素（表格/图片段落）向前查找最近的非空正文段落作为参考段落，
        用于让标记段落继承周围正文字体样式。

        Args:
            elements: 文档 body 子元素列表
            target_element: 目标元素（表格或图片段落）

        Returns:
            参考段落 XML 元素，若找不到则返回 None
        """
        from lxml import etree as _etree

        try:
            idx = elements.index(target_element)
        except ValueError:
            return None

        for search_idx in range(idx - 1, -1, -1):
            elem = elements[search_idx]
            local_name = _etree.QName(elem.tag).localname
            if local_name == 'p':
                text = self._get_paragraph_xml_text(elem)
                if text.strip():
                    return elem
            elif local_name == 'tbl':
                continue
        return None

    def _get_paragraph_xml_text(self, p_element) -> str:
        """提取段落 XML 元素的纯文本（拼接所有 w:t 节点）。"""
        t_elements = p_element.findall(f'.//{WNS_P}t')
        texts = [t.text for t in t_elements if t.text]
        return ''.join(texts)

    def _create_marker_paragraph(self, doc, text: str, ref_element=None):
        """
        创建标记段落：继承参考段落的字体样式并居中对齐。

        Args:
            doc: Word 文档对象
            text: 段落文本
            ref_element: 参考段落 XML 元素（用于继承字体样式）

        Returns:
            新创建的段落对象
        """
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from lxml import etree as _etree

        if ref_element is not None:
            try:
                new_p = copy.deepcopy(ref_element)

                first_run_with_style = None
                for run in new_p.findall(f'{WNS_P}r'):
                    rPr = run.find(f'{WNS_P}rPr')
                    if rPr is not None:
                        first_run_with_style = run
                        break

                if first_run_with_style is None:
                    first_run_with_style = new_p.find(f'{WNS_P}r')

                if first_run_with_style is not None:
                    rPr = first_run_with_style.find(f'{WNS_P}rPr')
                    for child in list(first_run_with_style):
                        first_run_with_style.remove(child)
                    if rPr is not None:
                        first_run_with_style.append(rPr)
                    t_elem = _etree.SubElement(first_run_with_style, f'{WNS_P}t')
                    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    t_elem.text = text

                    for child in list(new_p):
                        local_name = _etree.QName(child.tag).localname
                        if local_name == 'r' and child != first_run_with_style:
                            new_p.remove(child)
                        elif local_name not in ('r', 'pPr'):
                            new_p.remove(child)

                pPr = new_p.find(f'{WNS_P}pPr')
                if pPr is None:
                    pPr = _etree.SubElement(new_p, f'{WNS_P}pPr')
                jc = pPr.find(f'{WNS_P}jc')
                if jc is None:
                    jc = _etree.SubElement(pPr, f'{WNS_P}jc')
                jc.set(f'{WNS_P}val', 'center')

                from docx.text.paragraph import Paragraph
                return Paragraph(new_p, doc)

            except Exception as e:
                logger.debug(f"复制段落样式失败，使用默认样式: {e}")

        p = doc.add_paragraph()
        p.text = text
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return p

    def _check_paragraph_has_image(self, p_element: 'etree._Element') -> bool:
        """
        判断段落是否包含需要标记的图片元素。

        检测标准图片（blip / drawing / pict），以及图像类 OLE 嵌入对象
        （Visio、PowerPoint、公式、艺术字等），排除 Excel 等表格类对象。

        Args:
            p_element: 段落 XML 元素 (w:p)

        Returns:
            是否包含需要标记的图片
        """
        for elem in p_element.iter():
            tag_local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

            if tag_local in ('blip', 'drawing', 'pict'):
                return True

            if tag_local == 'object':
                for ole_elem in elem.iter():
                    ole_tag = ole_elem.tag.split('}')[-1] if '}' in ole_elem.tag else ole_elem.tag
                    if ole_tag == 'OLEObject':
                        prog_id = ole_elem.get('ProgID', '')
                        for image_prog_prefix in self._IMAGE_OLE_PROGIDS:
                            if prog_id.startswith(image_prog_prefix):
                                return True
                        break

        return False

    def _process_nested_table_contents(self, doc: 'Document', tbl_element: 'etree._Element',
                                        parent_table_idx: int,
                                        counters: dict) -> None:
        """
        递归处理表格单元格内的嵌套内容，在嵌套表格与图片前后插入标记段落。

        Args:
            doc: Document 对象
            tbl_element: 表格 XML 元素 (w:tbl)
            parent_table_idx: 父表格编号
            counters: 计数器字典，格式 {'nested_table': {parent_idx: count}, 'cell_image': {parent_idx: count}}
        """
        if parent_table_idx not in counters['nested_table']:
            counters['nested_table'][parent_table_idx] = 0
        if parent_table_idx not in counters['cell_image']:
            counters['cell_image'][parent_table_idx] = 0

        for tr in tbl_element:
            if not tr.tag.endswith('tr'):
                continue
            for tc in tr:
                if not tc.tag.endswith('tc'):
                    continue

                insertions = []

                for idx, child in enumerate(list(tc)):
                    if child.tag.endswith('tbl'):
                        counters['nested_table'][parent_table_idx] += 1
                        nested_idx = counters['nested_table'][parent_table_idx]
                        marker_name = f"Table_{parent_table_idx}_{nested_idx}"

                        start_marker = self._create_marker_paragraph_element(
                            f"{{{{{marker_name}_Start}}}}"
                        )
                        end_marker = self._create_marker_paragraph_element(
                            f"{{{{{marker_name}_End}}}}"
                        )

                        insertions.append((idx, 'before', start_marker))
                        insertions.append((idx + 1, 'after', end_marker, child))

                        self._process_nested_table_contents(doc, child, nested_idx, counters)

                    elif child.tag.endswith('p'):
                        if self._check_paragraph_has_image(child):
                            counters['cell_image'][parent_table_idx] += 1
                            image_idx = counters['cell_image'][parent_table_idx]
                            marker_name = f"Image_{parent_table_idx}_{image_idx}"

                            start_marker = self._create_marker_paragraph_element(f"{{{{{marker_name}_Start}}}}")
                            end_marker = self._create_marker_paragraph_element(f"{{{{{marker_name}_End}}}}")

                            insertions.append((idx, 'before', start_marker))
                            insertions.append((idx + 1, 'after', end_marker, child))

                for insertion in reversed(insertions):
                    if len(insertion) == 3:
                        idx, pos, marker = insertion
                        if pos == 'before':
                            tc.insert(idx, marker)
                    elif len(insertion) == 4:
                        idx, pos, marker, _ = insertion
                        if pos == 'after':
                            tc.insert(idx, marker)

    def mark_tables_and_images(self, word_path: Path, work_dir: Path) -> Optional[Path]:
        """
        标记 Word 文档中的表格与图片。

        复制原文档后遍历 body 元素，在表格、图片段落前后插入标记段落，
        并递归处理表格单元格内的嵌套表格与图片。已存在标记文件时直接复用。

        Returns:
            标记后的 Word 文件路径；失败返回 None
        """
        try:
            marked_file_path = work_dir / f"{word_path.stem}_marked{word_path.suffix}"

            if marked_file_path.exists():
                logger.info(f"标记文件已存在，跳过: {marked_file_path} ")
                return marked_file_path

            try:
                from docx import Document  # type: ignore
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
                import shutil

                marked_file_path.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(word_path, marked_file_path)

                doc = Document(marked_file_path)
                table_count = 0
                image_count = 0

                nested_counters = {
                    'nested_table': {},
                    'cell_image': {},
                }

                body = doc._element.body
                elements = list(body)
                new_elements = []

                for element in elements:
                    if element.tag.endswith('tbl'):
                        table_count += 1
                        ref_p = self._find_ref_paragraph_before(elements, element)
                        start_p = self._create_marker_paragraph(
                            doc, f"{{{{Table_{table_count}_Start}}}}", ref_p
                        )
                        new_elements.append(start_p._element)
                        new_elements.append(element)
                        end_p = self._create_marker_paragraph(
                            doc, f"{{{{Table_{table_count}_End}}}}", ref_p
                        )
                        new_elements.append(end_p._element)

                        self._process_nested_table_contents(doc, element, table_count, nested_counters)

                    elif element.tag.endswith('p'):
                        has_image = self._check_paragraph_has_image(element)
                        if has_image:
                            image_count += 1
                            ref_p = self._find_ref_paragraph_before(elements, element)
                            start_p = self._create_marker_paragraph(
                                doc, f"{{{{Image_{image_count}_Start}}}}", ref_p
                            )
                            new_elements.append(start_p._element)
                            new_elements.append(element)
                            end_p = self._create_marker_paragraph(
                                doc, f"{{{{Image_{image_count}_End}}}}", ref_p
                            )
                            new_elements.append(end_p._element)
                        else:
                            new_elements.append(element)
                    else:
                        new_elements.append(element)

                body.clear()
                for elem in new_elements:
                    body.append(elem)

                doc.save(marked_file_path)

                total_nested_tables = sum(nested_counters['nested_table'].values())
                total_cell_images = sum(nested_counters['cell_image'].values())
                logger.info(f"Word标记成功（python-docx）: {table_count}表格, {image_count}顶层图片, "
                           f"{total_nested_tables}嵌套表格, {total_cell_images}单元格内图片")
                return marked_file_path

            except Exception as e:
                logger.error(f"python-docx标记失败: {e}", exc_info=True)
                return None

        except Exception as e:
            logger.error(f"Word标记失败: {e}", exc_info=True)
            return None


word_marker = WordRegionMarker()


class WordRegionExtractor:
    """
    Word区域提取器

    从包含 {{Table_N_Start}} / {{Table_N_End}} 或
    {{Image_N_Start}} / {{Image_N_End}} 标记的Word文档中
    提取每个区域，生成独立的、格式完整的Word文件。

    核心思路：
      - 以源文档为模板整体复制（保留所有样式、主题、字体、关系）
      - 直接操作 document.xml，删除不属于目标区域的元素
      - 重新打包为合法的 .docx，无需手动处理图片关系
    """

    # 匹配开始/结束标记，如 {{Table_1_Start}} 或 {{TemplateTable_1_Start}}
    START_PATTERN = re.compile(r'^\{\{(TemplateTable|Table|Image)_(\d+)(?:_(\d+))?_Start\}\}$')
    END_PATTERN = re.compile(r'^\{\{(TemplateTable|Table|Image)_(\d+)(?:_(\d+))?_End\}\}$')


    def extract_regions(self, marked_file: str, export_dir: str) -> List[Dict[str, str]]:
        """
        从标记的Word文档中提取所有区域，输出为独立Word文件。

        优化：源文档只解压一次，每个区域只打包其实际引用的媒体文件。

        Args:
            marked_file: 含标记的源 .docx 路径
            export_dir:  输出目录

        Returns:
            [{"name": "Table_1_Start", "path": "/some/dir/Table_1_Start.docx"}, ...]
        """
        source_path = Path(marked_file)
        export_path = Path(export_dir)
        export_path.mkdir(parents=True, exist_ok=True)

        # 用 python-docx 读取文档，仅用于解析区域边界索引
        doc = Document(str(source_path))
        regions = self._find_regions(doc)
        logger.info(f"共找到 {len(regions)} 个区域: {[r['name'] for r in regions]}")

        # ---- 初始化临时目录变量 ----
        tmp_dir = None
        
        try:
            # ---- 预解压源文档一次，解析关系映射 ----
            tmp_dir = tempfile.mkdtemp()
            logger.info(f"临时目录创建在: {tmp_dir}")
            
            extract_dir = Path(tmp_dir) / 'extracted'
            with zipfile.ZipFile(source_path, 'r') as zf:
                zf.extractall(extract_dir)
            logger.info(f"源文档已解压到临时目录: {extract_dir}")

            rels_map = self._parse_relationships(extract_dir)
            common_files = self._classify_source_files(extract_dir)
            logger.info(f"预分类: {len(common_files)} 个公共文件")

            # ---- 预解析 document.xml（一次解析，避免每个区域重复） ----
            doc_xml_path = extract_dir / 'word' / 'document.xml'
            doc_tree = etree.parse(
                str(doc_xml_path), etree.XMLParser(remove_blank_text=False)
            )
            doc_root = doc_tree.getroot()
            doc_body = doc_tree.find(f'.//{WNS_P}body')
            if doc_body is None:
                raise ValueError("document.xml 中找不到 <w:body>")
            doc_sect_pr = doc_body.find(f'{WNS_P}sectPr')
            # body 子元素列表（段落+表格），含索引，与 _find_regions 对齐
            body_children = self._get_indexed_content_elements(list(doc_body))
            # 提取根元素命名空间声明，后续构建 document.xml 时需要
            nsmap = doc_root.nsmap
            logger.info(
                f"document.xml 预解析完成，body子元素: {len(body_children)}, "
                f"命名空间: {list(nsmap.keys())}"
            )

            results = []
            for region in regions:
                output_name = f"{region['name']}_Start"
                output_file = export_path / f"{output_name}.docx"
                try:
                    self._export_region_by_clone(
                        extract_dir, rels_map, common_files,
                        doc_root, body_children, doc_sect_pr,
                        nsmap, region, output_file,
                    )
                    results.append({"name": output_name, "path": str(output_file)})
                    logger.info(f"成功导出: {output_file.name}")
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    logger.error(f"导出区域 {region['name']} 失败: {e}", exc_info=True)

            return results
        
        except Exception as e:
            logger.error(f"处理过程中发生错误: {e}", exc_info=True)
            raise  # 或者 return [] 根据需求决定
        
        finally:
            # ---- 安全清理临时目录 ----
            if tmp_dir is not None:
                tmp_path = Path(tmp_dir)
                if tmp_path.exists():
                    try:
                        shutil.rmtree(tmp_path, ignore_errors=True)
                        logger.debug(f"临时目录已清理: {tmp_dir}")
                    except PermissionError as e:
                        logger.warning(f"清理临时目录时权限错误: {e}")
                    except Exception as e:
                        logger.warning(f"清理临时目录时发生错误: {e}")


    def _find_regions(self, doc: Document) -> List[Dict]:
        """
        遍历文档体，收集所有区域的 (name, start_idx, end_idx)。

        索引基于 body 的直接子元素列表（段落 + 表格），
        start_idx 指向开始标记的下一个元素，
        end_idx   指向结束标记本身（不含）。
        """
        # 收集 body 直接子元素及类型
        elements = self._collect_body_elements(doc)

        # 递归查找所有标记段落（包括嵌套在表格单元格内的）
        all_markers = self._find_all_markers_recursive(doc._element.body)

        regions: List[Dict] = []
        # 使用栈结构处理嵌套区域
        region_stack: List[Dict] = []

        for marker_info in all_markers:
            text = marker_info['text']
            elem = marker_info['element']
            is_in_cell = marker_info['is_in_cell']
            cell_path = marker_info['cell_path']

            start_match = self.START_PATTERN.match(text)
            if start_match:
                marker_type = start_match.group(1)
                primary_idx = int(start_match.group(2))
                secondary_idx = int(start_match.group(3)) if start_match.group(3) else None

                # 判断是否为嵌套区域：通过是否有第二个索引判断
                # Table_1 或 Image_1 是顶层，Table_1_1 或 Image_1_1 是嵌套
                is_nested = secondary_idx is not None

                if is_nested:
                    # 嵌套区域使用元素引用
                    name = f"{marker_type}_{primary_idx}_{secondary_idx}"
                    new_region = {
                        'name':       name,
                        'type':       marker_type.lower(),
                        'start_idx':  None,  # 嵌套区域不使用索引
                        'end_idx':    None,
                        'is_nested':  True,
                        'start_elem': elem,
                        'end_elem':   None,
                        'cell_path':  cell_path,
                    }
                else:
                    # 顶层区域使用索引
                    name = f"{marker_type}_{primary_idx}"
                    # 查找标记段落在 elements 列表中的索引
                    idx = self._find_element_index(elements, elem)
                    new_region = {
                        'name':       name,
                        'type':       marker_type.lower(),
                        'start_idx':  idx + 1,   # 内容从下一个元素开始
                        'end_idx':    None,
                        'is_nested':  False,
                        'start_elem': elem,
                        'end_elem':   None,
                    }

                # Push 到栈
                region_stack.append(new_region)
                continue

            end_match = self.END_PATTERN.match(text)
            if end_match and region_stack:
                marker_type = end_match.group(1)
                primary_idx = int(end_match.group(2))
                secondary_idx = int(end_match.group(3)) if end_match.group(3) else None

                # 检查栈顶是否匹配
                top_region = region_stack[-1]

                if top_region['is_nested']:
                    expected_name = f"{marker_type}_{primary_idx}_{secondary_idx}"
                else:
                    expected_name = f"{marker_type}_{primary_idx}"

                if top_region['name'] == expected_name:
                    # 匹配成功，Pop 并完成区域
                    top_region['end_elem'] = elem
                    if not top_region['is_nested']:
                        # 顶层区域需要计算结束索引
                        idx = self._find_element_index(elements, elem)
                        top_region['end_idx'] = idx   # 不含结束标记本身

                    region_stack.pop()
                    regions.append(top_region)
                else:
                    logger.warning(
                        f"结束标记 {expected_name} 与栈顶区域 {top_region['name']} 不匹配，忽略"
                    )

        # 检查未闭合的区域
        for unfinished in region_stack:
            logger.warning(f"区域 {unfinished['name']} 有开始标记但没有结束标记，已忽略")

        return regions

    def _find_element_index(self, elements: List[Tuple[str, etree._Element]],
                             target_elem: etree._Element) -> int:
        """在元素列表中查找目标元素的索引"""
        for idx, (elem_type, elem) in enumerate(elements):
            if elem is target_elem:
                return idx
        return -1

    def _find_all_markers_recursive(self, root_element: etree._Element) -> List[Dict]:
        """
        递归遍历文档树，查找所有标记段落

        包括嵌套在表格单元格内的标记段落

        Args:
            root_element: 根元素（通常是 body）

        Returns:
            标记信息列表，每个包含: text, element, is_in_cell, cell_path
        """
        markers: List[Dict] = []

        def traverse(elem: etree._Element, in_cell: bool = False, cell_path: List = None):
            """递归遍历元素树"""
            if cell_path is None:
                cell_path = []

            local_name = etree.QName(elem.tag).localname

            # 如果是段落，检查是否为标记段落
            if local_name == 'p':
                text = self._get_element_text(elem)
                if self.START_PATTERN.match(text) or self.END_PATTERN.match(text):
                    markers.append({
                        'text':       text,
                        'element':    elem,
                        'is_in_cell': in_cell,
                        'cell_path':  cell_path.copy(),
                    })

            # 如果是表格单元格，标记进入单元格
            elif local_name == 'tc':
                new_path = cell_path + [elem]
                for child in elem:
                    traverse(child, in_cell=True, cell_path=new_path)

            # 其他元素继续递归遍历
            else:
                for child in elem:
                    traverse(child, in_cell=in_cell, cell_path=cell_path)

        traverse(root_element)
        return markers

    def _export_region_by_clone(
        self,
        extract_dir: Path,
        rels_map: Dict[str, Tuple[str, str]],
        common_files: List[Tuple[str, Path]],
        doc_root: etree._Element,
        body_children: List[Tuple[int, etree._Element]],
        doc_sect_pr: Optional[etree._Element],
        nsmap: Dict,
        region: Dict,
        output_file: Path,
    ) -> None:
        """
        从预解压、预解析的源文档中裁剪区域并打包为最小化 docx。

        使用预解析的 document.xml 树，仅 deep-copy 目标元素，避免重复解析。
        """
        # 1. 从预解析树构建裁剪后的 document.xml
        if region.get('is_nested', False):
            trimmed_xml = self._build_nested_region_xml(
                doc_root, doc_sect_pr, nsmap, region,
            )
        else:
            trimmed_xml = self._build_top_level_region_xml(
                doc_root, body_children, doc_sect_pr, nsmap, region,
            )

        # 2. 扫描裁剪后XML中引用的 rId（用于保留 image 等内容性关系）
        referenced_rids = self._collect_referenced_rids(trimmed_xml)

        # 3. 构建过滤后的 rels 文件
        filtered_rels = self._build_filtered_rels(extract_dir, referenced_rids)

        # 4. 解析引用的实际文件路径
        referenced_files: Set[str] = set()
        for rid in referenced_rids:
            if rid in rels_map:
                target, _ = rels_map[rid]
                referenced_files.add(f"word/{target}")

        # 5. 创建最小化 docx
        output_file.parent.mkdir(parents=True, exist_ok=True)
        written: Set[str] = set()
        with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as zout:
            for arcname, file_path in common_files:
                zout.write(file_path, arcname)
                written.add(arcname)

            zout.writestr('word/document.xml', trimmed_xml)
            written.add('word/document.xml')
            zout.writestr('word/_rels/document.xml.rels', filtered_rels)
            written.add('word/_rels/document.xml.rels')

            for arcname in sorted(referenced_files):
                if arcname in written:
                    continue
                file_path = extract_dir / arcname
                if file_path.is_file():
                    ext = Path(arcname).suffix.lower()
                    compress = zipfile.ZIP_STORED if ext in _MEDIA_EXTENSIONS else zipfile.ZIP_DEFLATED
                    info = zipfile.ZipInfo(arcname)
                    info.compress_type = compress
                    with open(file_path, 'rb') as f:
                        zout.writestr(info, f.read())
                    written.add(arcname)

    # ---- 基于预解析树的快速 XML 构建（避免重复解析整个文件） ----

    def _build_top_level_region_xml(
        self,
        doc_root: etree._Element,
        body_children: List[Tuple[int, etree._Element]],
        doc_sect_pr: Optional[etree._Element],
        nsmap: Dict,
        region: Dict,
    ) -> bytes:
        """
        从预解析树构建顶层区域的 document.xml，仅 deep-copy 目标元素。
        """
        keep_start = region['start_idx']
        keep_end = region['end_idx']

        new_root = etree.Element(doc_root.tag, nsmap=nsmap)
        for key, value in doc_root.attrib.items():
            new_root.set(key, value)
        new_body = etree.SubElement(new_root, f'{WNS_P}body')

        for idx, elem in body_children:
            if keep_start <= idx < keep_end:
                copied = copy.deepcopy(elem)
                self._remove_internal_markers(copied)
                new_body.append(copied)

        # sectPr 放最后，删除页眉页脚引用
        if doc_sect_pr is not None:
            new_sect_pr = copy.deepcopy(doc_sect_pr)
            for ref in new_sect_pr.findall(f'{WNS_P}headerReference'):
                new_sect_pr.remove(ref)
            for ref in new_sect_pr.findall(f'{WNS_P}footerReference'):
                new_sect_pr.remove(ref)
            new_body.append(new_sect_pr)

        return etree.tostring(
            new_root, xml_declaration=True, encoding='UTF-8', standalone=True,
        )

    def _build_nested_region_xml(
        self,
        doc_root: etree._Element,
        doc_sect_pr: Optional[etree._Element],
        nsmap: Dict,
        region: Dict,
    ) -> bytes:
        """
        从预解析树构建嵌套区域的 document.xml，仅 deep-copy 目标元素。
        """
        start_elem = region.get('start_elem')
        end_elem = region.get('end_elem')
        cell_path = region.get('cell_path', [])
        region_type = region.get('type', '')

        if start_elem is None or end_elem is None:
            raise ValueError(f"嵌套区域 {region['name']} 缺少标记元素引用")

        container = cell_path[-1] if cell_path else None
        if container is None:
            raise ValueError(f"嵌套区域 {region['name']} 缺少单元格路径")

        # 获取标记之间的内容元素
        content_elements = self._get_elements_between_markers(
            start_elem, end_elem, container,
        )

        new_root = etree.Element(doc_root.tag, nsmap=nsmap)
        for key, value in doc_root.attrib.items():
            new_root.set(key, value)

        new_body = etree.SubElement(new_root, f'{WNS_P}body')

        if region_type == 'table':
            for elem in content_elements:
                tbl_elem = copy.deepcopy(elem)
                self._ensure_tbl_grid(tbl_elem)
                new_body.append(tbl_elem)
        else:
            simple_tbl = self._create_simple_table_for_nested(content_elements)
            new_body.append(simple_tbl)

        if doc_sect_pr is not None:
            new_sect_pr = copy.deepcopy(doc_sect_pr)
            for ref in new_sect_pr.findall(f'{WNS_P}headerReference'):
                new_sect_pr.remove(ref)
            for ref in new_sect_pr.findall(f'{WNS_P}footerReference'):
                new_sect_pr.remove(ref)
            new_body.append(new_sect_pr)

        return etree.tostring(
            new_root, xml_declaration=True, encoding='UTF-8', standalone=True,
        )

    @staticmethod
    def _classify_source_files(extract_dir: Path) -> List[Tuple[str, Path]]:
        """
        预分类解压目录中的文件，返回公共文件列表（一次性，避免每区域重复 rglob）。

        排除：document.xml、rels 文件、media 目录、页眉页脚。
        这些文件每区域单独处理。
        """
        common: List[Tuple[str, Path]] = []
        for file_path in sorted(extract_dir.rglob('*')):
            if not file_path.is_file():
                continue
            arcname = str(file_path.relative_to(extract_dir)).replace('\\', '/')

            if arcname == 'word/document.xml':
                continue
            if arcname == 'word/_rels/document.xml.rels':
                continue
            if arcname.startswith('word/media/'):
                continue
            if arcname.startswith('word/header') or arcname.startswith('word/footer'):
                continue

            common.append((arcname, file_path))
        return common

    # ---- 按需打包辅助方法 ----

    @staticmethod
    def _parse_relationships(extract_dir: Path) -> Dict[str, Tuple[str, str]]:
        """
        解析 word/_rels/document.xml.rels，返回 {rId: (target, type)} 映射。
        """
        rels_path = extract_dir / 'word' / '_rels' / 'document.xml.rels'
        if not rels_path.exists():
            return {}

        tree = etree.parse(str(rels_path))
        rels_map: Dict[str, Tuple[str, str]] = {}
        for rel in tree.getroot():
            rid = rel.get('Id')
            target = rel.get('Target', '')
            rel_type = rel.get('Type', '')
            if rid:
                rels_map[rid] = (target, rel_type)
        return rels_map

    @staticmethod
    def _collect_referenced_rids(xml_bytes: bytes) -> Set[str]:
        """
        扫描裁剪后的 document.xml，收集所有被引用的 rId。

        匹配 r:embed（图片/图表等嵌入资源）和 r:link（超链接）。
        """
        referenced: Set[str] = set()
        for pattern in (rb'r:embed="(rId\d+)"', rb'r:link="(rId\d+)"', rb'r:id="(rId\d+)"'):
            for match in re.findall(pattern, xml_bytes):
                referenced.add(match.decode('utf-8'))
        return referenced

    @staticmethod
    def _build_filtered_rels(
        extract_dir: Path,
        referenced_rids: Set[str],
    ) -> bytes:
        """
        构建过滤后的 document.xml.rels。

        保留策略：基于关系 Type 判断，而非仅按 rId 引用过滤。
        - 结构性关系（styles/settings/theme/numbering/fontTable/webSettings/
          footnotes/endnotes/customXml 等）一律保留
        - image 关系仅保留被 document.xml 内容引用的
        - header/footer 关系一律移除（sectPr 中已删除引用）
        """
        rels_path = extract_dir / 'word' / '_rels' / 'document.xml.rels'
        ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
        empty_rels = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<Relationships xmlns="{ns}"></Relationships>'
        )
        if not rels_path.exists():
            return empty_rels.encode('utf-8')

        tree = etree.parse(str(rels_path))
        root = tree.getroot()

        # OOXML 标准关系类型后缀 — 这些是文档结构性关系，必须保留
        _STRUCTURAL_SUFFIXES = (
            '/styles',
            '/settings',
            '/theme',
            '/numbering',
            '/fontTable',
            '/webSettings',
            '/footnotes',
            '/endnotes',
            '/customXml',
            '/comments',
            '/commentsExtended',
            '/commentsExtensible',
            '/commentsIds',
            '/people',
        )

        for rel in list(root):
            rel_type = rel.get('Type', '')

            # 1. header/footer 关系 → 移除
            if '/header' in rel_type or '/footer' in rel_type:
                root.remove(rel)
                continue

            # 2. image 关系 → 仅保留被引用的
            if '/image' in rel_type:
                rid = rel.get('Id')
                if rid not in referenced_rids:
                    root.remove(rel)
                continue

            # 3. 结构性关系 → 保留
            if any(rel_type.endswith(suffix) for suffix in _STRUCTURAL_SUFFIXES):
                continue

            # 4. 其他类型（如超链接等）→ 仅保留被引用的
            rid = rel.get('Id')
            if rid not in referenced_rids:
                root.remove(rel)

        return etree.tostring(
            tree,
            xml_declaration=True,
            encoding='UTF-8',
            standalone=True,
        )

    def _remove_internal_markers(self, elem: etree._Element) -> None:
        """
        清理元素内部的标记段落

        删除表格单元格内的 {{Table_N_M_Start/End}} 和 {{Image_N_M_Start/End}} 标记

        Args:
            elem: 要清理的元素（通常是表格）
        """
        # 只处理表格元素
        if not elem.tag.endswith('tbl'):
            return

        # 遍历表格结构：tbl -> tr -> tc -> p
        for tr in elem:
            if not tr.tag.endswith('tr'):
                continue
            for tc in tr:
                if not tc.tag.endswith('tc'):
                    continue

                # 收集需要删除的标记段落
                markers_to_remove = []
                for child in tc:
                    if child.tag.endswith('p'):
                        text = self._get_element_text(child)
                        # 匹配嵌套标记：{{Table_N_M_Start/End}} 或 {{Image_N_M_Start/End}}
                        if self._is_internal_marker(text):
                            markers_to_remove.append(child)

                # 删除标记段落
                for marker in markers_to_remove:
                    tc.remove(marker)

    def _is_internal_marker(self, text: str) -> bool:
        """
        判断文本是否为内部标记

        内部标记格式：{{Table_N_M_Start/End}} 或 {{Image_N_M_Start/End}}
        特点：有两个数字索引（N_M）

        Args:
            text: 段落文本

        Returns:
            是否为内部标记
        """
        if not text.startswith('{{') or not text.endswith('}}'):
            return False

        # 移除 {{ 和 }}
        content = text[2:-2]

        # 检查格式：Table_N_M_Start/End 或 Image_N_M_Start/End
        parts = content.split('_')
        if len(parts) != 4:
            return False

        type_part, num1, num2, start_end = parts
        if type_part not in ('Table', 'Image'):
            return False

        if start_end not in ('Start', 'End'):
            return False

        try:
            int(num1)
            int(num2)
            return True
        except ValueError:
            return False

    def _ensure_tbl_grid(self, tbl_elem: etree._Element) -> None:
        """
        确保表格元素有必需的 tblGrid 子元素

        Word 要求每个表格都有 tblGrid 元素来定义列宽

        Args:
            tbl_elem: 表格 XML 元素
        """
        # 检查是否已有 tblGrid
        tblGrid = tbl_elem.find(f'{WNS_P}tblGrid')
        if tblGrid is not None:
            return  # 已存在，无需添加

        # 计算列数（通过第一行的单元格数）
        col_count = 0
        for tr in tbl_elem:
            if etree.QName(tr.tag).localname == 'tr':
                for tc in tr:
                    if etree.QName(tc.tag).localname == 'tc':
                        col_count += 1
                break  # 只检查第一行

        if col_count == 0:
            col_count = 1  # 默认至少1列

        # 创建 tblGrid 元素
        tblGrid = etree.Element(f'{WNS_P}tblGrid')
        for _ in range(col_count):
            gridCol = etree.SubElement(tblGrid, f'{WNS_P}gridCol')
            gridCol.set(f'{WNS_P}w', '3000')  # 默认列宽

        # 找到合适的位置插入 tblGrid
        # tblGrid 应该在 tblPr 之后，如果 tblPr 存在
        tblPr = tbl_elem.find(f'{WNS_P}tblPr')
        if tblPr is not None:
            # 在 tblPr 之后插入
            tblPr_index = list(tbl_elem).index(tblPr)
            tbl_elem.insert(tblPr_index + 1, tblGrid)
        else:
            # 在第一个位置插入 tblPr 和 tblGrid
            new_tblPr = etree.Element(f'{WNS_P}tblPr')
            tbl_elem.insert(0, new_tblPr)
            tbl_elem.insert(1, tblGrid)


    def _get_elements_between_markers(self, start_marker: etree._Element,
                                       end_marker: etree._Element,
                                       container: Optional[etree._Element]) -> List[etree._Element]:
        """
        获取开始标记和结束标记之间的内容元素

        Args:
            start_marker: 开始标记段落元素
            end_marker: 结束标记段落元素
            container: 包含这些标记的容器元素（通常是单元格）

        Returns:
            内容元素列表（不含标记段落本身）
        """
        content_elements: List[etree._Element] = []

        if container is not None:
            # 在容器内查找
            found_start = False
            for child in container:
                if child is start_marker:
                    found_start = True
                    continue
                if child is end_marker:
                    break
                if found_start:
                    content_elements.append(child)

        return content_elements

    def _create_simple_table_for_nested(self, content_elements: List[etree._Element]) -> etree._Element:
        """
        为嵌套内容创建简化的表格结构

        创建一个单行单列的表格，包含给定的内容元素
        使用深拷贝的方式保留原有元素的命名空间

        Args:
            content_elements: 要包含的内容元素列表

        Returns:
            表格 XML 元素 (w:tbl)
        """
        import copy

        # 创建表格元素（使用命名空间）
        tbl = etree.Element(f'{WNS_P}tbl', nsmap={None: WNS})

        # 创建表格属性
        tblPr = etree.SubElement(tbl, f'{WNS_P}tblPr')

        # 创建 tblGrid（必需元素，定义列宽）
        tblGrid = etree.SubElement(tbl, f'{WNS_P}tblGrid')
        gridCol = etree.SubElement(tblGrid, f'{WNS_P}gridCol')
        gridCol.set(f'{WNS_P}w', '5000')  # 默认列宽

        # 创建表格行
        tr = etree.SubElement(tbl, f'{WNS_P}tr')

        # 创建表格单元格
        tc = etree.SubElement(tr, f'{WNS_P}tc')

        # 添加内容元素到单元格（深拷贝以保留完整结构）
        for elem in content_elements:
            # 深拷贝元素，保留所有子元素和属性
            tc.append(copy.deepcopy(elem))

        return tbl

    def _get_indexed_content_elements(
        self,
        all_children: List[etree._Element],
    ) -> List[Tuple[int, etree._Element]]:
        """
        从 body 子元素列表中提取段落和表格，并给出连续索引。
        索引与 _find_regions 中的索引保持一致。
        """
        result: List[Tuple[int, etree._Element]] = []
        idx = 0
        for elem in all_children:
            local = etree.QName(elem.tag).localname
            if local in ('p', 'tbl'):
                result.append((idx, elem))
                idx += 1
        return result

    @staticmethod
    def _collect_body_elements(doc: Document) -> List[Tuple[str, etree._Element]]:
        """收集文档 body 的直接子元素（段落 + 表格）"""
        elements: List[Tuple[str, etree._Element]] = []
        for elem in doc.element.body:
            local = etree.QName(elem.tag).localname
            if local == 'p':
                elements.append(('paragraph', elem))
            elif local == 'tbl':
                elements.append(('table', elem))
        return elements

    @staticmethod
    def _get_element_text(p_element: etree._Element) -> str:
        """提取段落元素的纯文本（拼接所有 w:t 子节点）

        注意：不使用 itertext()，因为它在某些情况下会返回重复文本。
        改用 findall 精确查找所有 w:t 元素。
        """
        try:
            # 使用 findall 精确查找所有 w:t 元素，避免重复
            t_elements = p_element.findall(f'.//{WNS_P}t')
            texts = [t.text for t in t_elements if t.text]
            return ''.join(texts).strip()
        except Exception:
            return ''

word_region_extractor = WordRegionExtractor()