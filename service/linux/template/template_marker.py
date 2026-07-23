"""
模板标记器

功能说明：
- 标记Word模板中的表格
- 获取表格标题（表格前一段落文本）
- 添加{{Table_N_Start/End}}标记和{{标题}}段落

主要类：
- TemplateMarker: 模板标记器类
"""

# ========== 标准库导入 ==========
import logging
import re
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Dict, List, Tuple, Optional

# ========== 第三方库导入 ==========
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

# ========== 模块配置 ==========
logger = logging.getLogger(__name__)

# Word XML 命名空间
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WNS_P = f'{{{WNS}}}'


class TemplateMarker:
    """
    模板标记器

    标记Word文档中的表格，并获取表格标题。
    标题规则：
    - 默认取表格前一段落的文本
    - 截断到50字符
    - 如果前一段落是空段落或标记段落，标题为"无标题表格"
    """

    # 标记模式匹配
    MARKER_PATTERN = re.compile(r'^\{\{(TemplateTable|Image)_(\d+)_Start\}\}$')
    END_MARKER_PATTERN = re.compile(r'^\{\{(TemplateTable|Image)_(\d+)_End\}\}$')

    # 标题最大长度
    MAX_TITLE_LENGTH = 20

    # 默认标题
    DEFAULT_TITLE = "无标题表格"

    def mark_tables_with_titles(
        self,
        word_path: Path,
        work_dir: Path
    ) -> Tuple[Path, List[Dict]]:
        """
        标记表格并获取标题

        Args:
            word_path: 原Word文件路径
            work_dir: 工作目录

        Returns:
            (标记后的文件路径, 表格信息列表)
            表格信息: [{"index": 1, "title": "表格标题"}, ...]
        """
        try:
            # 标记文件路径
            marked_file_path = work_dir / f"{word_path.stem}_marked{word_path.suffix}"

            # 如果已存在标记文件，直接返回（需要重新解析标题）
            if marked_file_path.exists():
                logger.info(f"标记文件已存在，跳过标记: {marked_file_path}")
                table_titles = self._parse_table_titles_from_marked(marked_file_path)
                return marked_file_path, table_titles

            # 复制原文件
            shutil.copy2(word_path, marked_file_path)

            # 加载文档
            doc = Document(marked_file_path)

            # 获取文档body元素
            body = doc._element.body
            elements = list(body)

            # 收集表格信息
            table_titles: List[Dict] = []
            table_count = 0

            # 新元素列表（重建文档结构）
            new_elements = []

            # 遍历元素，记录表格位置和标题
            table_positions = []  # [(index_in_elements, table_element), ...]

            for idx, element in enumerate(elements):
                local_name = etree.QName(element.tag).localname
                if local_name == 'tbl':
                    table_count += 1
                    table_positions.append((idx, element, table_count))

            # 重新遍历，插入标记
            for idx, element in enumerate(elements):
                local_name = etree.QName(element.tag).localname

                if local_name == 'tbl':
                    # 找到对应的表格编号
                    table_info = None
                    for pos_idx, pos_elem, pos_num in table_positions:
                        if pos_idx == idx:
                            table_info = (pos_num, pos_elem)
                            break

                    if table_info:
                        table_num, tbl_element = table_info

                        # 获取标题（表格前一段落文本）
                        title = self._get_table_title(elements, idx)

                        # 截断标题
                        title = self._truncate_title(title)

                        # 记录表格信息
                        table_titles.append({
                            "index": table_num,
                            "title": title
                        })

                        # 获取参考段落（表格内部文字）用于继承样式
                        ref_p_element = self._get_ref_paragraph_from_table(tbl_element)

                        # 插入标题段落（继承参考段落样式）
                        title_p = self._create_marker_paragraph(
                            doc, f"{{{{标题：{title}}}}}", ref_p_element, center=True
                        )
                        new_elements.append(title_p._element)

                        # 插入开始标记（继承参考段落样式）
                        start_p = self._create_marker_paragraph(
                            doc, f"{{{{TemplateTable_{table_num}_Start}}}}", ref_p_element, center=True
                        )
                        new_elements.append(start_p._element)

                        # 表格本身
                        new_elements.append(element)

                        # 插入结束标记（继承参考段落样式）
                        end_p = self._create_marker_paragraph(
                            doc, f"{{{{TemplateTable_{table_num}_End}}}}", ref_p_element, center=True
                        )
                        new_elements.append(end_p._element)

                        logger.info(f"表格 {table_num}: 标题='{title}'")

                elif local_name == 'p':
                    # 段落：检查是否包含图片
                    has_image = any(
                        child.tag.endswith(('blip', 'drawing', 'pict'))
                        for child in element.iter()
                    )
                    if has_image:
                        # 图片段落：暂不处理，直接保留
                        new_elements.append(element)
                    else:
                        new_elements.append(element)

                else:
                    # 其他元素直接保留
                    new_elements.append(element)

            # 重建文档body
            body.clear()
            for elem in new_elements:
                body.append(elem)

            # 保存标记文件
            doc.save(marked_file_path)

            logger.info(f"标记完成: {table_count} 个表格, 输出: {marked_file_path}")

            return marked_file_path, table_titles

        except Exception as e:
            import traceback
            traceback.print_exc()
            logger.error(f"标记表格失败: {e}", exc_info=True)
            raise

    def _get_ref_paragraph_from_table(self, tbl_element):
        """
        从表格内部获取参考段落用于继承样式

        Args:
            tbl_element: 表格XML元素

        Returns:
            参考段落XML元素，如果没有则返回 None
        """
        try:
            # 表格结构: tbl -> tr -> tc -> p
            # 查找第一个有文字内容的段落
            for tr in tbl_element.findall(f'{WNS_P}tr'):
                for tc in tr.findall(f'{WNS_P}tc'):
                    for p in tc.findall(f'{WNS_P}p'):
                        # 检查段落是否有文字内容
                        text = self._get_paragraph_text(p)
                        if text.strip():
                            return p

            # 如果没有找到有内容的段落，返回第一个单元格的第一个段落
            first_tr = tbl_element.find(f'{WNS_P}tr')
            if first_tr is not None:
                first_tc = first_tr.find(f'{WNS_P}tc')
                if first_tc is not None:
                    first_p = first_tc.find(f'{WNS_P}p')
                    if first_p is not None:
                        return first_p

        except Exception as e:
            logger.debug(f"从表格获取参考段落失败: {e}")

        # 没有找到有效段落
        return None

    def _create_marker_paragraph(
        self,
        doc: Document,
        text: str,
        ref_element=None,
        center: bool = True
    ):
        """
        创建标记段落，继承参考段落的字体样式

        Args:
            doc: Word文档对象
            text: 段落文本
            ref_element: 参考段落XML元素（用于继承样式）
            center: 是否居中对齐

        Returns:
            新创建的段落对象
        """
        # 尝试从参考段落复制样式
        if ref_element is not None:
            try:
                # 深拷贝参考段落
                new_p = deepcopy(ref_element)

                # 查找第一个有 rPr 的 run（包含字体样式）
                first_run_with_style = None
                for run in new_p.findall(f'{WNS_P}r'):
                    rPr = run.find(f'{WNS_P}rPr')
                    if rPr is not None:
                        first_run_with_style = run
                        break

                # 如果没有带样式的 run，用第一个 run
                if first_run_with_style is None:
                    first_run_with_style = new_p.find(f'{WNS_P}r')

                if first_run_with_style is not None:
                    # 保留 rPr，清空其他内容，设置新文本
                    rPr = first_run_with_style.find(f'{WNS_P}rPr')

                    # 清空 run 中的所有子元素
                    for child in list(first_run_with_style):
                        first_run_with_style.remove(child)

                    # 重新添加 rPr（如果有）
                    if rPr is not None:
                        first_run_with_style.append(rPr)

                    # 添加新的文本元素
                    t_elem = etree.SubElement(first_run_with_style, f'{WNS_P}t')
                    # 保留空格
                    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    t_elem.text = text

                    # 清空段落中其他 run，只保留这一个
                    for child in list(new_p):
                        local_name = etree.QName(child.tag).localname
                        if local_name == 'r' and child != first_run_with_style:
                            new_p.remove(child)
                        elif local_name not in ('r', 'pPr'):
                            new_p.remove(child)

                # 设置居中对齐（覆盖原对齐方式）
                if center:
                    pPr = new_p.find(f'{WNS_P}pPr')
                    if pPr is None:
                        pPr = etree.SubElement(new_p, f'{WNS_P}pPr')
                    jc = pPr.find(f'{WNS_P}jc')
                    if jc is None:
                        jc = etree.SubElement(pPr, f'{WNS_P}jc')
                    jc.set(f'{WNS_P}val', 'center')

                # 创建段落对象包装
                from docx.text.paragraph import Paragraph
                return Paragraph(new_p, doc)

            except Exception as e:
                logger.debug(f"复制段落样式失败，使用默认样式: {e}")

        # 回退：创建默认段落
        p = doc.add_paragraph()
        p.text = text
        if center:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return p

    def _get_table_title(self, elements: List, table_idx: int) -> str:
        """
        获取表格标题（表格前一段落文本）

        Args:
            elements: 文档元素列表
            table_idx: 表格在元素列表中的索引

        Returns:
            表格标题文本
        """
        # 从表格前一个元素开始向前查找
        for search_idx in range(table_idx - 1, -1, -1):
            element = elements[search_idx]
            local_name = etree.QName(element.tag).localname

            if local_name == 'p':
                # 获取段落文本
                text = self._get_paragraph_text(element)

                # 跳过空段落
                if not text.strip():
                    continue

                # 跳过标记段落
                if self.MARKER_PATTERN.match(text.strip()) or \
                   self.END_MARKER_PATTERN.match(text.strip()):
                    continue

                # 找到有效标题
                return text.strip()

            elif local_name == 'tbl':
                # 前一个元素是表格，跳过
                continue

        # 没有找到有效标题
        return self.DEFAULT_TITLE

    def _get_paragraph_text(self, p_element) -> str:
        """
        获取段落的纯文本

        Args:
            p_element: 段落XML元素

        Returns:
            段落文本
        """
        try:
            # 查找所有 w:t 元素
            t_elements = p_element.findall(f'.//{WNS_P}t')
            texts = [t.text for t in t_elements if t.text]
            return ''.join(texts)
        except Exception:
            return ''

    def _truncate_title(self, title: str) -> str:
        """
        截断标题到最大长度

        Args:
            title: 原标题

        Returns:
            截断后的标题
        """
        if len(title) <= self.MAX_TITLE_LENGTH:
            return title
        return title[:self.MAX_TITLE_LENGTH] + "..."

    def _parse_table_titles_from_marked(self, marked_file: Path) -> List[Dict]:
        """
        从已标记文件中解析表格标题

        用于处理已存在的标记文件，提取标题信息

        Args:
            marked_file: 已标记的Word文件

        Returns:
            表格信息列表 [{"index": 1, "title": "xxx"}, ...]
        """
        try:
            doc = Document(marked_file)

            table_titles: List[Dict] = []
            table_count = 0

            # 模式匹配标题段落 {{标题}}
            title_pattern = re.compile(r'^\{\{(.+?)\}\}$')

            for element in doc._element.body:
                local_name = etree.QName(element.tag).localname

                if local_name == 'p':
                    text = self._get_paragraph_text(element).strip()

                    # 检查是否是标题段落
                    title_match = title_pattern.match(text)
                    if title_match:
                        title = title_match.group(1)
                        # 跳过标记段落
                        if title.startswith('TemplateTable_') or title.startswith('Image_'):
                            continue
                        # 这是表格标题，下一个应该是TemplateTable_N_Start
                        table_count += 1
                        table_titles.append({
                            "index": table_count,
                            "title": title
                        })

                elif local_name == 'tbl':
                    # 表格元素，计数已在标题段落处理
                    pass

            logger.info(f"从标记文件解析: {len(table_titles)} 个表格标题")
            return table_titles

        except Exception as e:
            import traceback
            traceback.print_exc()
            logger.warning(f"解析标记文件标题失败: {e}")
            return []


# ========== 全局单例 ==========
_template_marker: Optional[TemplateMarker] = None


def get_template_marker() -> TemplateMarker:
    """获取模板标记器单例"""
    global _template_marker
    if _template_marker is None:
        _template_marker = TemplateMarker()
    return _template_marker